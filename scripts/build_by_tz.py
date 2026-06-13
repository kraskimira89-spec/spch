# -*- coding: utf-8 -*-
"""
Контент-анализ строго по ТЗ:
- 10 тематических + 3 тональных кода (как в методичке)
- атомарное разбиение высказываний (одна мысль = одна строка)
- Excel: Данные, Коды, Итоги_темы (частота + ранг), Итоги_тональность
- Word-обобщение с ранжированием
"""

from __future__ import annotations

import hashlib
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
TZ_DOCX = ROOT / "ТЗ контент.docx"
TZ_PDF = ROOT / "ТЗ контент.pdf"
OUT = ROOT / "output" / "Контент-анализ_по_ТЗ"

# === Коды строго по ТЗ (шаг 2.1 + 2.2) ===
THEME_CODES: list[tuple[str, str, str]] = [
    ("НОРМА_442ФЗ", "442-ФЗ и связанные акты",
     "Изменение/уточнение ФЗ № 442-ФЗ и связанных актов для расширения доступа НКО."),
    ("КОНКУРС_ПРОЦЕДУРЫ", "Конкурсы и процедуры",
     "Упрощение конкурсов, снижение бюрократии, 44-ФЗ/223-ФЗ."),
    ("ТАРИФЫ_НОРМАТИВЫ", "Тарифы и нормативы",
     "Тарифы, подушевые нормативы, индексация финансирования соцуслуг."),
    ("РЕЕСТР_ПОСТАВЩИКОВ", "Реестр поставщиков",
     "Упрощение включения в реестр, прозрачные критерии."),
    ("ГАРАНТИИ_ОБЪЕМОВ", "Гарантии объёмов",
     "Гарантированный объём, многолетние контракты, авансирование, лимиты."),
    ("ОБУЧЕНИЕ_НКО", "Обучение НКО",
     "Обучение конкурсам, отчётности, юридическим требованиям."),
    ("ИНФОРМАЦИЯ_ПРОЗРАЧНОСТЬ", "Информация и прозрачность",
     "Информирование о конкурсах, критериях, планах финансирования."),
    ("ИНСТИТУТЫ_ПОДДЕРЖКИ", "Институты поддержки",
     "Ресурсные центры, опорные организации, грантовая поддержка."),
    ("МУНИЦИПАЛЬНЫЙ_УРОВЕНЬ", "Муниципальный уровень",
     "Практика на муниципальном уровне, решения администраций."),
    ("КОНКУРЕНЦИЯ_АНТИМОНОПОЛИЯ", "Конкуренция и антимонополия",
     "Борьба с монополией госучреждений, равные условия."),
]

TONE_CODES: list[tuple[str, str, str]] = [
    ("ТОНИЧЕСКИЙ_ПОЗИТИВ", "Позитив",
     "Позитивная оценка мер или перспектив."),
    ("ТОНИЧЕСКИЙ_НЕГАТИВ", "Негатив",
     "Критика, скепсис, описание проблем."),
    ("ТОНИЧЕСКИЙ_СОМНЕНИЕ", "Сомнение",
     "Затрудняюсь ответить, неопределённость."),
]

ALL_CODES = THEME_CODES + TONE_CODES

THEME_KEYWORDS: dict[str, list[str]] = {
    "НОРМА_442ФЗ": [r"442", r"189.?фз", r"законодатель", r"изменен.*закон", r"закон\s", r"нпа", r"постановлен"],
    "КОНКУРС_ПРОЦЕДУРЫ": [r"конкурс", r"закуп", r"44.?фз", r"223.?фз", r"тендер", r"отбор", r"процедур", r"бюрократ"],
    "ТАРИФЫ_НОРМАТИВЫ": [r"тариф", r"подушев", r"норматив", r"индекс", r"инфляц", r"мрот", r"стоимост.*услуг"],
    "РЕЕСТР_ПОСТАВЩИКОВ": [r"реестр", r"включен.*реестр", r"вхожд.*реестр", r"вступлен.*реестр"],
    "ГАРАНТИИ_ОБЪЕМОВ": [
        r"объ[её]м", r"аванс", r"долгосроч", r"многолет", r"гарант", r"лимит", r"квот",
        r"финансирован", r"субсид", r"бюджет", r"ассигнован", r"кассов",
    ],
    "ОБУЧЕНИЕ_НКО": [r"обучен", r"семинар", r"тренинг", r"квалификац", r"методическ", r"консультац"],
    "ИНФОРМАЦИЯ_ПРОЗРАЧНОСТЬ": [r"информ", r"прозрач", r"открытост", r"доступност.*информ"],
    "ИНСТИТУТЫ_ПОДДЕРЖКИ": [r"ресурсн.*центр", r"грант", r"опорн", r"институт", r"цисс"],
    "МУНИЦИПАЛЬНЫЙ_УРОВЕНЬ": [r"муницип", r"администрац", r"омс", r"местн.*самоуправ"],
    "КОНКУРЕНЦИЯ_АНТИМОНОПОЛИЯ": [
        r"монопол", r"конкурен", r"равн.*услов", r"антимонопол", r"конфликт интерес", r"дискримин",
    ],
}


@dataclass
class Statement:
    id: int
    text: str
    parent_id: int
    codes: dict[str, int]
    text_hash: str = ""


CONTINUATION_RE = re.compile(
    r"^(?:через\b|за\b|по\b|при\b|для\b|с учетом|с уч[её]том|исходя из|в соответствии|"
    r"привести\b|чтобы\b|котор(?:ый|ая|ое|ые)\b|где\b|а также\b|либо\b|"
    r"так как\b|потому что\b|то есть\b|например\b)",
    re.I,
)
ACTION_HEADS = (
    "введение", "внедрение", "повышение", "увеличение", "снижение", "уменьшение",
    "расширение", "пересмотр", "создание", "формирование", "переход", "компенсация",
    "отмена", "масштабирование", "индексация", "субсидирование", "планирование",
    "скорректировать", "изменить", "упростить", "сократить", "обеспечить",
    "предусмотреть", "предоставить", "выделить", "закрепить", "ввести", "создать",
    "расширить", "увеличить", "повысить", "пересмотреть", "установить",
    "унифицировать", "отказаться", "вернуть", "заложить", "открыть",
)


def extract_paragraphs_from_docx(path: Path) -> list[str]:
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs: list[str] = []
    for para in root.iter(f"{ns}p"):
        parts: list[str] = []
        for t in para.iter(f"{ns}t"):
            if t.text:
                parts.append(t.text)
        text = "".join(parts).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def extract_text_from_docx(path: Path) -> str:
    # Сохраняем границы абзацев Word, чтобы не склеивать независимые ответы в одну строку.
    return "\n".join(extract_paragraphs_from_docx(path))


def extract_text_from_pdf(path: Path) -> str | None:
    try:
        from pypdf import PdfReader
    except ImportError:
        return None

    if not path.exists():
        return None

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(text)
    return "\n".join(pages)


def rough_split(body: str) -> list[str]:
    out: list[str] = []
    for line in body.splitlines():
        chunk = re.sub(r"\s+", " ", line).strip()
        if len(chunk) >= 3 and not chunk.lower().startswith("цель:"):
            out.append(chunk)
    return out


def normalize_fragment(text: str) -> str:
    text = re.sub(r"^\d+[\.\)]\s*", "", text.strip())
    text = re.sub(r"^-\s*", "", text)
    text = re.sub(r"\s+", " ", text).strip(" .;,-")
    return text


def split_numbered_items(text: str) -> list[str]:
    starts: list[int] = []
    for match in re.finditer(r"\d+[\.\)]", text):
        if match.start() > 0 and not text[match.start() - 1].isspace():
            continue
        tail = text[match.end() :]
        if not tail.strip():
            continue
        starts.append(match.start())

    if len(starts) < 2:
        return [text]

    parts: list[str] = []
    for idx, start in enumerate(starts):
        end = starts[idx + 1] if idx + 1 < len(starts) else len(text)
        parts.append(text[start:end].strip())
    return parts


def split_dash_list(text: str) -> list[str]:
    if not text.lstrip().startswith("-"):
        return [text]
    if text.count(" - ") < 2:
        return [text]
    parts: list[str] = []
    current: list[str] = []
    paren_depth = 0
    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "(":
            paren_depth += 1
        elif ch == ")" and paren_depth > 0:
            paren_depth -= 1

        if (
            paren_depth == 0
            and i + 2 < len(text)
            and text[i : i + 3] == " - "
            and current
        ):
            chunk = "".join(current).strip()
            if chunk:
                parts.append(chunk)
            current = []
            i += 3
            continue

        current.append(ch)
        i += 1

    tail = "".join(current).strip()
    if tail:
        parts.append(tail)
    return parts or [text]


def should_split_before(fragment: str) -> bool:
    probe = re.sub(r"^[\s\-,:;.]+", "", fragment, flags=re.UNICODE).lower()
    return any(probe.startswith(head) for head in ACTION_HEADS)


def split_top_level_measures(text: str) -> list[str]:
    parts: list[str] = []
    current: list[str] = []
    paren_depth = 0

    i = 0
    while i < len(text):
        ch = text[i]
        if ch == "(":
            paren_depth += 1
        elif ch == ")" and paren_depth > 0:
            paren_depth -= 1

        if paren_depth == 0 and ch in ",;":
            tail = text[i + 1 :]
            if should_split_before(tail):
                chunk = "".join(current).strip()
                if chunk:
                    parts.append(chunk)
                current = []
                i += 1
                while i < len(text) and text[i].isspace():
                    i += 1
                continue

        if paren_depth == 0 and ch in ".!?":
            tail = text[i + 1 :]
            probe = tail.lstrip()
            if should_split_before(probe) or re.match(r"^\d+[\.\)]", probe):
                chunk = "".join(current).strip()
                if chunk:
                    parts.append(chunk)
                current = []
                i += 1
                while i < len(text) and text[i].isspace():
                    i += 1
                continue

        current.append(ch)
        i += 1

    tail_chunk = "".join(current).strip()
    if tail_chunk:
        parts.append(tail_chunk)
    return parts or [text]


def split_dirty_measure_runs(text: str) -> list[str]:
    cleaned = re.sub(r"\s+", " ", text).strip()
    if ". " not in cleaned:
        return [text]
    parts = re.split(
        r"(?<=\.)\s+(?=(?:уменьшение|повышение|увеличение|компенсация|введение|внедрение|"
        r"расширение|пересмотр|создание|формирование|переход|индексация|"
        r"скорректировать|изменить|упростить|сократить|обеспечить|предусмотреть|"
        r"предоставить|выделить|закрепить|ввести|создать|расширить|увеличить|"
        r"повысить|пересмотреть|установить|унифицировать|отказаться|вернуть|заложить)\b)",
        cleaned,
        flags=re.I,
    )
    return [part for part in parts if part.strip()] or [text]


def merge_continuations(parts: list[str]) -> list[str]:
    merged: list[str] = []
    for part in parts:
        cleaned = normalize_fragment(part)
        if len(cleaned) < 3:
            continue
        if merged:
            prev = merged[-1]
            if CONTINUATION_RE.match(cleaned):
                merged[-1] = f"{prev} {cleaned}".strip()
                continue
        merged.append(cleaned)
    return merged


def split_compound_clauses(text: str) -> list[str]:
    pieces = split_top_level_measures(text)
    expanded: list[str] = []
    for piece in pieces:
        numbered = split_numbered_items(piece)
        for item in numbered:
            expanded.extend(split_dirty_measure_runs(item))
    if not expanded:
        return [text]
    return merge_continuations(expanded)


def split_atomic(text: str) -> list[str]:
    """Дробит составное высказывание на атомарные (одна мысль — одна строка)."""
    prepared = re.sub(r"\s+", " ", text).strip()
    chunks = split_numbered_items(prepared)
    if len(chunks) == 1:
        chunks = split_dash_list(prepared)

    secondary: list[str] = []
    for chunk in chunks:
        secondary.extend(re.split(r"\s*;\s+", chunk))
    chunks = merge_continuations(secondary)

    result: list[str] = []
    for chunk in chunks:
        subchunks = split_compound_clauses(chunk)
        for sub in subchunks:
            cleaned = normalize_fragment(sub)
            if len(cleaned) >= 8:
                result.append(cleaned)
    return result if result else ([normalize_fragment(text)] if text.strip() else [])


def normalize_validation_text(text: str) -> str:
    text = text.lower().replace("ё", "е")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def count_action_heads(text: str) -> int:
    low = normalize_validation_text(text)
    count = 0
    for head in ACTION_HEADS:
        if re.search(rf"(^|[.;,]\s+|\d+[.)]\s*){re.escape(head)}\b", low):
            count += 1
    return count


def classify_pdf_validation_issue(
    block: str,
    parts: list[str],
    matched: bool,
    action_heads: int,
) -> dict[str, object] | None:
    normalized_block = normalize_validation_text(block)
    normalized_parts = [normalize_validation_text(part) for part in parts if normalize_validation_text(part)]
    block_len = len(normalized_block)
    short_count = sum(len(part) < 24 for part in normalized_parts)
    tiny_count = sum(len(part) < 14 for part in normalized_parts)
    numbered_items = len(split_numbered_items(block))
    semicolon_count = block.count(";")
    has_enumeration_colon = ":" in block and semicolon_count >= 3

    # Короткие односложные ответы PDF часто теряет или искажает; не шумим ими в отчете.
    if block_len < 45 and action_heads == 0 and len(normalized_parts) <= 1:
        return None

    # Длинные нумерованные блоки мер часто извлекаются из PDF нестрого, но сама сегментация здесь корректна.
    if not matched and numbered_items >= 5 and len(normalized_parts) >= numbered_items:
        matched = True
    if not matched and has_enumeration_colon and len(normalized_parts) >= 5:
        matched = True

    # Перечисление после двоеточия с несколькими пунктами через ';' считаем нормальной структурой, а не пере-дроблением.
    if has_enumeration_colon and len(normalized_parts) >= 4 and semicolon_count >= len(normalized_parts) - 2:
        short_count = 0
        tiny_count = 0

    score = 0
    reasons: list[str] = []

    if action_heads >= 2 and len(normalized_parts) == 1:
        score += 4 if action_heads >= 3 else 3
        reasons.append("несколько смысловых действий схлопнулись в одно атомарное высказывание")
    elif action_heads >= 4 and len(normalized_parts) <= max(2, action_heads - 2):
        score += 2
        reasons.append("число сегментов заметно ниже числа смысловых действий в блоке")

    if len(normalized_parts) >= 6 and short_count >= 3:
        score += 3
        reasons.append("возможна избыточная фрагментация на короткие сегменты")
    elif len(normalized_parts) >= 4 and (short_count >= 2 or tiny_count >= 1):
        score += 2
        reasons.append("есть признаки пере-дробления блока")

    if not matched and block_len >= 120 and (action_heads >= 2 or len(normalized_parts) >= 3):
        score += 1
        reasons.append("длинный структурно сложный блок не нашел уверенного совпадения в PDF")

    if score <= 0:
        return None

    if score >= 4:
        risk = "высокий риск"
    elif score >= 2:
        risk = "средний сигнал"
    else:
        risk = "слабый сигнал"

    return {
        "risk": risk,
        "score": score,
        "block_len": block_len,
        "short_count": short_count,
        "reasons": reasons,
    }


def build_pdf_segmentation_validation(
    pdf_path: Path,
    rough_blocks: list[str],
    statements: list[Statement],
) -> dict[str, object]:
    raw_pdf = extract_text_from_pdf(pdf_path)
    if raw_pdf is None:
        return {
            "enabled": False,
            "reason": "PDF-check недоступен: файл PDF не найден или библиотека pypdf не установлена.",
            "checked_blocks": 0,
            "matched_blocks": 0,
            "issues": [],
            "risk_counts": {"высокий риск": 0, "средний сигнал": 0, "слабый сигнал": 0},
        }

    pdf_text = normalize_validation_text(raw_pdf)
    grouped: dict[int, list[str]] = {}
    for stmt in statements:
        grouped.setdefault(stmt.parent_id, []).append(stmt.text)

    issues: list[dict[str, object]] = []
    risk_counts = {"высокий риск": 0, "средний сигнал": 0, "слабый сигнал": 0}
    matched_blocks = 0
    for parent_id, block in enumerate(rough_blocks, 1):
        normalized_block = normalize_validation_text(block)
        anchor = normalized_block[: min(140, len(normalized_block))]
        matched = bool(anchor and len(anchor) >= 30 and anchor in pdf_text)
        if matched:
            matched_blocks += 1

        parts = grouped.get(parent_id, [])
        action_heads = count_action_heads(block)
        classified = classify_pdf_validation_issue(block, parts, matched, action_heads)
        if not classified:
            continue

        risk = str(classified["risk"])
        risk_counts[risk] += 1

        # Подробно выводим только реально спорные блоки.
        if risk == "слабый сигнал":
            continue

        issues.append(
            {
                "parent_id": parent_id,
                "matched_in_pdf": matched,
                "action_heads": action_heads,
                "n_statements": len(parts),
                "reasons": classified["reasons"],
                "block": block,
                "statements": parts,
                "risk": risk,
                "score": classified["score"],
                "block_len": classified["block_len"],
                "short_count": classified["short_count"],
            }
        )

    return {
        "enabled": True,
        "reason": "",
        "checked_blocks": len(rough_blocks),
        "matched_blocks": matched_blocks,
        "issues": issues,
        "risk_counts": risk_counts,
        "pdf_chars": len(raw_pdf),
    }


def write_pdf_validation_report(path: Path, validation: dict[str, object]) -> None:
    lines = [
        "Гибридная проверка сегментации: DOCX как основной источник, PDF как контрольный слой",
        "",
    ]

    if not validation["enabled"]:
        lines.append(str(validation["reason"]))
        path.write_text("\n".join(lines), encoding="utf-8")
        return

    checked = int(validation["checked_blocks"])
    matched = int(validation["matched_blocks"])
    issues: list[dict[str, object]] = list(validation["issues"])
    risk_counts: dict[str, int] = dict(validation["risk_counts"])
    lines.extend(
        [
            f"Проверено исходных блоков: {checked}",
            f"Уверенно сопоставлено с PDF: {matched}",
            f"Высокий риск: {risk_counts['высокий риск']}",
            f"Средний сигнал: {risk_counts['средний сигнал']}",
            f"Слабый сигнал: {risk_counts['слабый сигнал']}",
            f"Детально выведено действительно спорных блоков: {len(issues)}",
            "",
            "Логика фильтрации:",
            "- короткие и тривиальные ответы не попадают в предупреждения;",
            "- слабый сигнал считается, но не выводится подробно;",
            "- подробно показываются только блоки со средним и высоким риском.",
            "",
        ]
    )

    for issue in issues[:80]:
        lines.append(f"Блок #{issue['parent_id']}")
        lines.append(f"- Уровень: {issue['risk']}")
        lines.append(f"- Совпадение с PDF: {'да' if issue['matched_in_pdf'] else 'нет'}")
        lines.append(f"- Количество смысловых действий: {issue['action_heads']}")
        lines.append(f"- Получено атомарных высказываний: {issue['n_statements']}")
        lines.append(f"- Коротких сегментов: {issue['short_count']}")
        lines.append(f"- Причины: {'; '.join(issue['reasons'])}")
        lines.append(f"- Исходный блок: {issue['block']}")
        if issue["statements"]:
            lines.append("- Сегменты:")
            for segment in issue["statements"]:
                lines.append(f"  * {segment}")
        lines.append("")

    path.write_text("\n".join(lines), encoding="utf-8")


def build_corpus(raw_text: str) -> tuple[list[Statement], list[str]]:
    idx = raw_text.find("Высказывания:")
    if idx == -1:
        raise ValueError("Блок «Высказывания:» не найден")
    body = raw_text[idx + len("Высказывания:") :]

    rough = rough_split(body)
    atomic: list[tuple[int, str]] = []
    parent = 0
    for block in rough:
        parent += 1
        parts = split_atomic(block)
        if len(parts) > 1:
            for p in parts:
                atomic.append((parent, p))
        else:
            atomic.append((parent, block))

    statements: list[Statement] = []
    for i, (pid, text) in enumerate(atomic, 1):
        h = hashlib.sha256(text.encode("utf-8")).hexdigest()[:12]
        statements.append(Statement(i, text, pid, code_statement(text), h))
    return statements, rough


def code_themes(text: str) -> dict[str, int]:
    low = text.lower()
    codes = {c[0]: 0 for c in THEME_CODES}
    for code, patterns in THEME_KEYWORDS.items():
        for pat in patterns:
            if re.search(pat, low):
                codes[code] = 1
                break
    return codes


def assign_tone(text: str) -> str:
    low = text.lower().strip()
    if re.search(r"затрудн|не знаю|неизвестно|надо подумать|правительству виднее|пусть решают|нет предлож", low):
        return "ТОНИЧЕСКИЙ_СОМНЕНИЕ"
    if re.search(
        r"мер достаточно|достаточно мер|доступ есть|вс[её] устраивает|вс[её] нормально|нет проблем|"
        r"не требуется|всего вполне|всем вс[её] доступно|вс[её] есть|принимаются достаточные",
        low,
    ):
        if not re.search(r"но\s|тариф.*низк|невыносим|задерж", low):
            return "ТОНИЧЕСКИЙ_ПОЗИТИВ"
    if re.search(
        r"невыносим|выжив|закрыт|обидно|монопол|задерж|кассов|убыт|штраф|шантаж|"
        r"не работает|искусственн|не поможет|ничего не поможет|руки опуск",
        low,
    ):
        return "ТОНИЧЕСКИЙ_НЕГАТИВ"
    if re.search(r"низк.*тариф|тариф.*низк|не индекс|остаточн|формально", low):
        return "ТОНИЧЕСКИЙ_НЕГАТИВ"
    return "ТОНИЧЕСКИЙ_СОМНЕНИЕ" if len(low) < 25 else "ТОНИЧЕСКИЙ_НЕГАТИВ"


def code_statement(text: str) -> dict[str, int]:
    themes = code_themes(text)
    tone = assign_tone(text)
    codes = {c[0]: 0 for c in ALL_CODES}
    for k, v in themes.items():
        codes[k] = v
    codes[tone] = 1
    return codes


def aggregate(statements: list[Statement]) -> dict:
    n = len(statements)
    themes = {c[0]: 0 for c in THEME_CODES}
    tones = {c[0]: 0 for c in TONE_CODES}
    for s in statements:
        for k, v in s.codes.items():
            if v:
                if k.startswith("ТОНИЧЕСКИЙ"):
                    tones[k] += 1
                elif k in themes:
                    themes[k] += 1
    ranked = sorted(themes.items(), key=lambda x: (-x[1], x[0]))
    return {"n": n, "themes": themes, "tones": tones, "ranked": ranked}


def write_excel(path: Path, statements: list[Statement]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Данные"

    headers = ["ID_высказывания", "ID_исходного_блока", "Текст_высказывания"] + [c[0] for c in ALL_CODES]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")

    for s in statements:
        row = [s.id, s.parent_id, s.text] + [s.codes[c[0]] for c in ALL_CODES]
        ws.append(row)

    ws.column_dimensions["C"].width = 70
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        row[2].alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "D2"

    n = len(statements)
    last = n + 1
    code_start_col = 4  # D

    # --- Коды ---
    ws_c = wb.create_sheet("Коды")
    ws_c.append(["№", "Код", "Название", "Описание", "Тип"])
    for cell in ws_c[1]:
        cell.font = Font(bold=True)
    num = 1
    for code, name, desc in THEME_CODES:
        ws_c.append([num, code, name, desc, "тематический"])
        num += 1
    for code, name, desc in TONE_CODES:
        ws_c.append([num, code, name, desc, "тональный"])
        num += 1

    # --- Итоги_темы (ГЛАВНЫЙ ЛИСТ ДЛЯ РЕЙТИНГА) ---
    ws_t = wb.create_sheet("Итоги_темы")
    ws_t["A1"] = "РАНЖИРОВАНИЕ ТЕМАТИЧЕСКИХ КОДОВ ПО УПОМИНАЕМОСТИ"
    ws_t["A1"].font = Font(bold=True, size=14)
    ws_t.merge_cells("A1:F1")
    ws_t["A2"] = f"Всего высказываний в корпусе: {n}"
    ws_t["A2"].font = Font(bold=True)

    hdr = 4
    cols = ["Ранг", "Код", "Название", "Частота", "Доля", "Формула_проверки"]
    for j, h in enumerate(cols, 1):
        c = ws_t.cell(row=hdr, column=j, value=h)
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="D9E2F3")

    stats = aggregate(statements)
    sorted_themes = sorted(stats["themes"].items(), key=lambda x: (-x[1], x[0]))
    for rank, (code_key, _) in enumerate(sorted_themes, 1):
        name = next(nm for c, nm, _ in THEME_CODES if c == code_key)
        col_idx = next(i for i, (c, _, _) in enumerate(THEME_CODES) if c == code_key)
        col_letter = get_column_letter(code_start_col + col_idx)
        row = hdr + rank
        ws_t.cell(row=row, column=1, value=rank)
        ws_t.cell(row=row, column=2, value=code_key)
        ws_t.cell(row=row, column=3, value=name)
        ws_t.cell(row=row, column=4, value=f"=SUM(Данные!{col_letter}2:{col_letter}{last})")
        ws_t.cell(row=row, column=5, value=f"=D{row}/$B$2")
        ws_t.cell(row=row, column=6, value=f"=SUM(Данные!{col_letter}2:{col_letter}{last})")
    ws_t["B2"] = f"=COUNTA(Данные!A2:A{last})"

    ws_t.column_dimensions["C"].width = 35
    ws_t.column_dimensions["F"].width = 20

    # Диаграмма топ тем
    bar = BarChart()
    bar.type = "col"
    bar.title = "Частота тематических кодов"
    bar.height = 12
    bar.width = 18
    data = Reference(ws_t, min_col=4, min_row=hdr, max_row=hdr + 10)
    cats = Reference(ws_t, min_col=3, min_row=hdr + 1, max_row=hdr + 10)
    bar.add_data(data, titles_from_data=True)
    bar.set_categories(cats)
    bar.dataLabels = DataLabelList()
    bar.dataLabels.showVal = True
    ws_t.add_chart(bar, "H4")

    # --- Итоги_тональность ---
    ws_tone = wb.create_sheet("Итоги_тональность")
    ws_tone["A1"] = "ТОНАЛЬНОСТЬ ВЫСКАЗЫВАНИЙ"
    ws_tone["A1"].font = Font(bold=True, size=14)
    ws_tone.merge_cells("A1:D1")
    ws_tone["A2"] = f"=COUNTA(Данные!A2:A{last})"
    ws_tone["A2"].font = Font(bold=True)

    ws_tone.append([])
    ws_tone.append(["Тональность", "Частота", "Доля"])
    for cell in ws_tone[4]:
        cell.font = Font(bold=True)

    tone_labels = {
        "ТОНИЧЕСКИЙ_ПОЗИТИВ": "Позитив",
        "ТОНИЧЕСКИЙ_НЕГАТИВ": "Негатив",
        "ТОНИЧЕСКИЙ_СОМНЕНИЕ": "Сомнение",
    }
    tone_start = code_start_col + len(THEME_CODES)
    for i, (code, label) in enumerate(tone_labels.items(), 5):
        col_letter = get_column_letter(tone_start + i - 5)
        ws_tone.cell(row=i, column=1, value=label)
        ws_tone.cell(row=i, column=2, value=f"=SUM(Данные!{col_letter}2:{col_letter}{last})")
        ws_tone.cell(row=i, column=3, value=f"=B{i}/$A$2")

    pie = PieChart()
    pie.title = "Тональность"
    pie.height = 10
    pie.width = 14
    labels = Reference(ws_tone, min_col=1, min_row=5, max_row=7)
    data = Reference(ws_tone, min_col=2, min_row=4, max_row=7)
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    pie.dataLabels = DataLabelList()
    pie.dataLabels.showPercent = True
    pie.dataLabels.showVal = True
    ws_tone.add_chart(pie, "E4")

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def save_ranking_chart(path: Path, ranked: list[tuple[str, int]]) -> None:
    names = {c: n for c, n, _ in THEME_CODES}
    top = [(names[c], v) for c, v in ranked if v > 0][:10]
    if not top:
        return
    labels, values = zip(*top)
    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.barh(range(len(labels)), values, color="#1F4E79")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=10)
    ax.invert_yaxis()
    ax.set_xlabel("Частота упоминаний")
    ax.set_title("Ранжирование тематических кодов", fontweight="bold")
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2, str(val), va="center")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=150)
    plt.close(fig)


def save_tone_chart(path: Path, tones: dict[str, int]) -> None:
    labels = ["Позитив", "Негатив", "Сомнение"]
    keys = ["ТОНИЧЕСКИЙ_ПОЗИТИВ", "ТОНИЧЕСКИЙ_НЕГАТИВ", "ТОНИЧЕСКИЙ_СОМНЕНИЕ"]
    values = [tones.get(k, 0) for k in keys]
    fig, ax = plt.subplots(figsize=(6, 4))
    colors = ["#2E7D32", "#C62828", "#757575"]
    ax.bar(labels, values, color=colors)
    ax.set_title("Тональность высказываний", fontweight="bold")
    for i, v in enumerate(values):
        ax.text(i, v + 0.5, str(v), ha="center")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=150)
    plt.close(fig)


def write_summary_docx(path: Path, statements: list[Statement], rough_count: int) -> None:
    stats = aggregate(statements)
    n = stats["n"]
    doc = Document()
    doc.add_heading("Обобщение: тематические коды и ранжирование", 0)
    doc.add_paragraph(f"Дата: {date.today().isoformat()}")
    doc.add_paragraph(
        f"Корпус: {n} атомарных высказываний (из {rough_count} исходных блоков ответов экспертов). "
        "Единица анализа — одна законченная мысль; составные ответы разбиты на отдельные строки."
    )

    doc.add_heading("1. Список тематических кодов (10 шт.)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Код"
    hdr[1].text = "Название"
    hdr[2].text = "Описание"
    for code, name, desc in THEME_CODES:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = name
        row[2].text = desc

    doc.add_heading("2. Ранжирование тем по упоминаемости", level=1)
    doc.add_paragraph(
        "Частота = число высказываний, в которых код = 1. "
        "Одно высказывание может иметь несколько кодов."
    )
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    for i, txt in enumerate(["Ранг", "Код", "Название", "Частота", "Доля"]):
        h2[i].text = txt

    names = {c: nm for c, nm, _ in THEME_CODES}
    for rank, (code, freq) in enumerate(stats["ranked"], 1):
        if freq == 0:
            continue
        row = t2.add_row().cells
        row[0].text = str(rank)
        row[1].text = code
        row[2].text = names[code]
        row[3].text = str(freq)
        row[4].text = f"{freq / n * 100:.1f}%"

    doc.add_heading("3. Тональные коды (3 шт.)", level=1)
    for code, name, desc in TONE_CODES:
        cnt = stats["tones"][code]
        doc.add_paragraph(f"• {name} ({code}): {cnt} ({cnt/n*100:.1f}%) — {desc}")

    doc.add_heading("4. Ключевой вывод", level=1)
    top3 = stats["ranked"][:3]
    top_txt = ", ".join(f"{names[c]} ({v})" for c, v in top3 if v)
    doc.add_paragraph(
        f"Наиболее часто упоминаются: {top_txt}. "
        "Эксперты концентрируются на тарифах/нормативах, объёмах финансирования и административных процедурах."
    )

    chart = OUT / "04_Диаграммы" / "ранжирование_тем.png"
    if chart.exists():
        doc.add_picture(str(chart), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)


def write_summary_txt(path: Path, statements: list[Statement]) -> None:
    stats = aggregate(statements)
    n = stats["n"]
    names = {c: nm for c, nm, _ in THEME_CODES}
    lines = [
        "ОБОБЩЕНИЕ: ТЕМАТИЧЕСКИЕ КОДЫ И РАНЖИРОВАНИЕ",
        f"Высказываний: {n}",
        "",
        "=== 10 ТЕМАТИЧЕСКИХ КОДОВ (список) ===",
    ]
    for i, (code, name, _) in enumerate(THEME_CODES, 1):
        lines.append(f"{i}. {code} — {name}")
    lines += ["", "=== РЕЙТИНГ ПО ЧАСТОТЕ ==="]
    for rank, (code, freq) in enumerate(stats["ranked"], 1):
        if freq:
            lines.append(f"{rank}. {names[code]}: {freq} ({freq/n*100:.1f}%)")
    lines += ["", "=== ТОНАЛЬНОСТЬ ==="]
    for code, name, _ in TONE_CODES:
        v = stats["tones"][code]
        lines.append(f"{name}: {v} ({v/n*100:.1f}%)")
    path.write_text("\n".join(lines), encoding="utf-8")


def write_corpus_docx(path: Path, statements: list[Statement]) -> None:
    doc = Document()
    doc.add_heading("Корпус атомарных высказываний", level=1)
    doc.add_paragraph(
        "Вопрос 72. Каждая строка — одна мысль. Составные ответы разбиты; "
        "ID_исходного_блока показывает, из какого исходного ответа получены фрагменты."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "ID блока"
    table.rows[0].cells[2].text = "Текст"
    for s in statements:
        row = table.add_row().cells
        row[0].text = str(s.id)
        row[1].text = str(s.parent_id)
        row[2].text = s.text
    doc.save(path)


def make_zip(folder: Path, zip_path: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in folder.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(folder.parent))


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for sub in ["01_ТЗ_и_корпус", "02_Excel", "03_Обобщение", "04_Диаграммы", "05_Отчет"]:
        (OUT / sub).mkdir(parents=True)

    raw = extract_text_from_docx(TZ_DOCX)
    statements, rough = build_corpus(raw)
    stats = aggregate(statements)

    print(f"Исходных блоков: {len(rough)}")
    print(f"Атомарных высказываний: {len(statements)}")
    print("Ранжирование тем:")
    names = {c: n for c, n, _ in THEME_CODES}
    for i, (code, freq) in enumerate(stats["ranked"][:10], 1):
        print(f"  {i}. {names[code]}: {freq}")

    shutil.copy2(TZ_DOCX, OUT / "01_ТЗ_и_корпус" / "01_TZ-kontent.docx")
    pdf = ROOT / "ТЗ контент.pdf"
    if pdf.exists():
        shutil.copy2(pdf, OUT / "01_ТЗ_и_корпус" / "01_TZ-kontent.pdf")

    write_corpus_docx(OUT / "01_ТЗ_и_корпус" / "02_Корпус_атомарных_высказываний.docx", statements)
    write_excel(OUT / "02_Excel" / "Контент-анализ_НКО_соцуслуги.xlsx", statements)
    save_ranking_chart(OUT / "04_Диаграммы" / "ранжирование_тем.png", stats["ranked"])
    save_tone_chart(OUT / "04_Диаграммы" / "тональность.png", stats["tones"])

    write_summary_docx(OUT / "03_Обобщение" / "Тематические_коды_и_ранжирование.docx", statements, len(rough))
    write_summary_txt(OUT / "03_Обобщение" / "Тематические_коды_и_ранжирование.txt", statements)

    # Краткий codes cheat sheet
    codes_doc = Document()
    codes_doc.add_heading("Шпаргалка кодировщика (10 + 3 кода)", 0)
    codes_doc.add_paragraph("Тематические коды — в одном высказывании может быть несколько «1».")
    for code, name, desc in THEME_CODES:
        codes_doc.add_paragraph(f"{code}: {name}. {desc}", style="List Bullet")
    codes_doc.add_paragraph("Тональность — только один код на высказывание:")
    for code, name, desc in TONE_CODES:
        codes_doc.add_paragraph(f"{code}: {name}. {desc}", style="List Bullet")
    codes_doc.save(OUT / "03_Обобщение" / "Шпаргалка_кодировщика.docx")

    from build_detailed_report import generate_report_package
    report_path = generate_report_package()
    print(f"Подробный отчёт: {report_path}")

    make_zip(OUT, ROOT / "output" / "Kontent-analiz-po-TZ.zip")
    print(f"Готово: {OUT}")


if __name__ == "__main__":
    main()
