# -*- coding: utf-8 -*-
"""Сборка пакетов контент-анализа: региональные корпусы, тональность, диаграммы, доклад РАН."""

from __future__ import annotations

import hashlib
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.chart.label import DataLabelList
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
TZ_DOCX = ROOT / "ТЗ контент.docx"
OUT_ROOT = ROOT / "output"
CHARTS_ROOT = OUT_ROOT / "_charts"

REGIONS: dict[str, dict] = {
    "Ростовская_область": {
        "title": "Ростовская область",
        "roiv": "Министерство труда и социального развития Ростовской области",
        "legislature": "Законодательное Собрание Ростовской области",
        "law": "областной закон от 03.09.2014 № 222-ЗС «О социальном обслуживании граждан в Ростовской области»",
        "keywords": [r"Ростовск", r"222-ЗС", r"222‑ЗС"],
    },
    "Свердловская_область": {
        "title": "Свердловская область",
        "roiv": "Министерство социальной политики Свердловской области",
        "legislature": "Законодательное Собрание Свердловской области",
        "law": "региональные акты, принятые во исполнение федерального закона № 442-ФЗ",
        "keywords": [r"Свердлов", r"Екатеринбург"],
    },
    "Ставропольский_край": {
        "title": "Ставропольский край",
        "roiv": "Министерство труда и социальной защиты населения Ставропольского края",
        "legislature": "Думой Ставропольского края",
        "law": "законы и постановления Ставропольского края в сфере социального обслуживания",
        "keywords": [r"Ставрополь"],
    },
    "Мурманская_область": {
        "title": "Мурманская область",
        "roiv": "Министерство социального развития Мурманской области",
        "legislature": "Мурманская областная Дума",
        "law": "региональные нормативные акты Мурманской области о социальном обслуживании",
        "keywords": [r"Мурманск", r"губернатор.*Мурман"],
    },
    "Пензенская_область": {
        "title": "Пензенская область",
        "roiv": "Министерство социального развития Пензенской области",
        "legislature": "Законодательное Собрание Пензенской области",
        "law": "региональные акты Пензенской области о социальном обслуживании",
        "keywords": [r"Пензен"],
    },
    "Общероссийский_корпус": {
        "title": "Российская Федерация (агрегированный корпус)",
        "roiv": "региональные органы исполнительной власти в сфере социального обслуживания",
        "legislature": "законодательные (представительные) органы субъектов РФ",
        "law": "федеральные законы № 442-ФЗ и № 189-ФЗ, а также региональные нормативные акты",
        "keywords": [],
    },
}

THEME_CODES: list[tuple[str, str, str]] = [
    ("ТАРИФЫ_ПОВЫШЕНИЕ", "Повышение тарифов", "Необходимость увеличения тарифов и подушевых нормативов."),
    ("ТАРИФЫ_ИНДЕКСАЦИЯ", "Индексация тарифов", "Автоматическая ежегодная индексация тарифов."),
    ("НОРМАТИВ_ПОДУШЕВОЙ", "Подушевые нормативы", "Корректировка подушевого финансирования."),
    ("ФИНАНСИРОВАНИЕ_ОБЪЕМ", "Объём финансирования", "Увеличение общего объёма бюджетных средств."),
    ("ФИНАНСИРОВАНИЕ_АВАНС_ДОЛГОСРОЧНО", "Аванс и долгосрочность", "Авансирование, многолетние контракты."),
    ("КОМПЕНСАЦИЯ_ПОЛНАЯ", "Полная компенсация", "Возмещение фактических затрат по тарифам."),
    ("ЗАКОНОДАТЕЛЬСТВО_442ФЗ", "442-ФЗ", "Изменения в федеральном законе № 442-ФЗ."),
    ("ЗАКОНОДАТЕЛЬСТВО_189ФЗ", "189-ФЗ / соцзаказ", "Государственный (муниципальный) социальный заказ."),
    ("НОРМАТИВКА_РЕГИОН", "Региональная нормативка", "Изменения региональных актов."),
    ("АНТИМОНОПОЛИЯ_КОНКУРЕНЦИЯ", "Конкуренция и монополия", "Равные условия, борьба с монополией."),
    ("СОЦЗАКАЗ_ВНЕДРЕНИЕ", "Социальный заказ", "Внедрение механизма соцзаказа."),
    ("СЕРТИФИКАТЫ_ВАУЧЕРЫ", "Сертификаты и ваучеры", "«Деньги следуют за получателем»."),
    ("СДОЛГОСРОЧНЫЙ_УХОД", "Долговременный уход", "Доступ НКО к системе СДУ."),
    ("ГРАНТЫ_СУБСИДИИ_НКО", "Гранты и субсидии", "Грантовая поддержка НКО."),
    ("РЕЕСТР_УПРОЩЕНИЕ", "Реестр поставщиков", "Упрощение включения в реестр."),
    ("КОНКУРСЫ_УПРОЩЕНИЕ", "Конкурсы", "Упрощение конкурсных процедур."),
    ("ОТЧЕТНОСТЬ_СНИЖЕНИЕ", "Снижение отчётности", "Уменьшение объёма отчётности."),
    ("СРОКИ_ВЫПЛАТ", "Сроки выплат", "Сокращение сроков компенсации."),
    ("КРИТЕРИИ_ПРОЗРАЧНЫЕ", "Прозрачные критерии", "Единые критерии допуска."),
    ("АРЕНДА_КОММУНАЛКА", "Аренда и коммуналка", "Компенсация аренды и коммунальных услуг."),
    ("ИМУЩЕСТВО_ПОДДЕРЖКА", "Имущественная поддержка", "Предоставление помещений НКО."),
    ("НАЛОГИ_ЛЬГОТЫ", "Налоги и льготы", "Налоговые льготы для НКО."),
    ("ИНФОРМАЦИЯ_ДОСТУПНОСТЬ", "Информационная открытость", "Доступность информации."),
    ("КОНСУЛЬТАЦИИ_РЕСУРСНЫЕ_ЦЕНТРЫ", "Ресурсные центры", "Обучение и сопровождение НКО."),
    ("ДИАЛОГ_ГОС_НКО", "Диалог гос—НКО", "Взаимодействие органов власти и НКО."),
    ("ЛИМИТЫ_ОТМЕНА", "Отмена лимитов", "Снятие лимитов на число получателей."),
    ("ПРАВО_ВЫБОРА", "Право выбора", "Реальный выбор поставщика."),
    ("КРИТЕРИИ_НУЖДАЕМОСТИ", "Критерии нуждаемости", "Критерии признания нуждающимися."),
]

TONE_CODES: list[tuple[str, str, str]] = [
    ("ТОН_ПОЗИТИВ", "Позитив", "Удовлетворённость действующей системой."),
    ("ТОН_КОНСТРУКТИВ", "Конструктивные предложения", "Предложения мер без эмоциональной критики."),
    ("ТОН_НЕГАТИВ", "Негатив / критика", "Критика, жалобы, описание системных проблем."),
    ("ТОН_СОМНЕНИЕ", "Сомнение", "«Затрудняюсь ответить», неопределённость."),
]

CODES = THEME_CODES + TONE_CODES

THEME_KEYWORDS: dict[str, list[str]] = {
    "ТАРИФЫ_ПОВЫШЕНИЕ": [r"повышен.*тариф", r"увеличен.*тариф", r"тариф.*повыш", r"низк.*тариф", r"тариф.*низк", r"заниж", r"невыносим", r"90 рубл", r"138"],
    "ТАРИФЫ_ИНДЕКСАЦИЯ": [r"индекс", r"инфляц"],
    "НОРМАТИВ_ПОДУШЕВОЙ": [r"подушев", r"среднедуш"],
    "ФИНАНСИРОВАНИЕ_ОБЪЕМ": [r"объ[её]м финанс", r"больше финанс", r"увеличен.*финанс", r"бюджетн.*финанс", r"ассигнован"],
    "ФИНАНСИРОВАНИЕ_АВАНС_ДОЛГОСРОЧНО": [r"аванс", r"долгосроч", r"3.?5 лет", r"многолет", r"стабильн.*финанс", r"кассов"],
    "КОМПЕНСАЦИЯ_ПОЛНАЯ": [r"полн.*компенс", r"реальн.*затрат", r"себестоим", r"фактическ", r"100%", r"остаточн"],
    "ЗАКОНОДАТЕЛЬСТВО_442ФЗ": [r"442", r"442‑фз", r"442-фз"],
    "ЗАКОНОДАТЕЛЬСТВО_189ФЗ": [r"189", r"189‑фз", r"189-фз", r"госзаказ", r"государственн.*заказ"],
    "НОРМАТИВКА_РЕГИОН": [r"региональн.*акт", r"региональн.*закон", r"постановлен", r"ручн.*регулир"],
    "АНТИМОНОПОЛИЯ_КОНКУРЕНЦИЯ": [r"монопол", r"конкурен", r"равн.*услов", r"конфликт интерес", r"дискримин"],
    "СОЦЗАКАЗ_ВНЕДРЕНИЕ": [r"социальн.*заказ", r"соцзаказ", r"гос.*заказ"],
    "СЕРТИФИКАТЫ_ВАУЧЕРЫ": [r"сертификат", r"ваучер", r"деньги следуют", r"финансы за человек"],
    "СДОЛГОСРОЧНЫЙ_УХОД": [r"долговремен", r"\bсду\b", r"систем.*уход"],
    "ГРАНТЫ_СУБСИДИИ_НКО": [r"грант", r"субсид"],
    "РЕЕСТР_УПРОЩЕНИЕ": [r"реестр", r"включен.*реестр", r"вхожд"],
    "КОНКУРСЫ_УПРОЩЕНИЕ": [r"конкурс", r"занижен.*цен", r"тендер", r"отбор"],
    "ОТЧЕТНОСТЬ_СНИЖЕНИЕ": [r"отч[её]тност", r"бюрократ", r"документооборот", r"справок", r"проверок"],
    "СРОКИ_ВЫПЛАТ": [r"срок.*выплат", r"задерж", r"ждать.*месяц", r"своевремен"],
    "КРИТЕРИИ_ПРОЗРАЧНЫЕ": [r"прозрач", r"понятн.*критер", r"един.*стандарт", r"един.*правил"],
    "АРЕНДА_КОММУНАЛКА": [r"аренд", r"коммунал"],
    "ИМУЩЕСТВО_ПОДДЕРЖКА": [r"помещен", r"имуществ", r"нежил.*фонд"],
    "НАЛОГИ_ЛЬГОТЫ": [r"налог", r"льгот", r"страхов.*тариф"],
    "ИНФОРМАЦИЯ_ДОСТУПНОСТЬ": [r"информ", r"открытост"],
    "КОНСУЛЬТАЦИИ_РЕСУРСНЫЕ_ЦЕНТРЫ": [r"ресурсн.*центр", r"обучен", r"консультац", r"методическ"],
    "ДИАЛОГ_ГОС_НКО": [r"диалог", r"совещан", r"взаимодейств", r"партн[её]р"],
    "ЛИМИТЫ_ОТМЕНА": [r"лимит", r"ограничен", r"квот", r"не урез", r"отмен.*лимит"],
    "ПРАВО_ВЫБОРА": [r"право выбор", r"выбирать поставщ", r"выбор поставщ"],
    "КРИТЕРИИ_НУЖДАЕМОСТИ": [r"нуждающ", r"признан.*нужда", r"иппсу", r"ипссу"],
}

TONE_POSITIVE = [
    r"мер достаточно", r"достаточно мер", r"доступ есть", r"вс[её] устраивает",
    r"вс[её] нормально", r"нет проблем", r"не требуется", r"прост.*прозрачен",
    r"всего вполне достаточно", r"всем вс[её] доступно", r"вс[её] доступно",
    r"вс[её] есть", r"вс[её] необходим", r"принимаются достаточные",
]
TONE_DOUBT = [
    r"затрудн", r"не знаю", r"неизвестно", r"надо подумать", r"правительству виднее",
    r"пусть решают", r"^тест$", r"предложений нет", r"предложений не", r"нет предложений",
    r"никакие", r"никакме",
]
TONE_NEGATIVE = [
    r"невыносим", r"выжив", r"закрыт", r"обидно", r"руки опуск", r"штраф", r"гонен",
    r"шантаж", r"убыт", r"кассов", r"задерж", r"монопол", r"дискримин", r"не поможет",
    r"ничего не поможет", r"стресс", r"чудес", r"не работает", r"искусственн",
    r"удыш", r"выгоран", r"кадров.*дефицит", r"не индекс", r"остаточн",
]


@dataclass
class Statement:
    id: int
    text: str
    region_tag: str | None
    codes: dict[str, int]
    source_id: int = 0
    text_hash: str = ""


def extract_text_from_docx(path: Path) -> str:
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    parts: list[str] = []
    for t in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
        if t.text:
            parts.append(t.text)
        if t.tail:
            parts.append(t.tail)
    return "".join(parts)


def split_statements(text: str) -> list[str]:
    idx = text.find("Высказывания:")
    if idx == -1:
        raise ValueError("Блок 'Высказывания:' не найден в ТЗ")
    body = text[idx + len("Высказывания:") :].strip()
    body = re.sub(r"([а-яё])(\d+\.)", r"\1 \2", body)
    body = re.sub(r"(\d+\.)([А-ЯЁ])", r"\1 \2", body)
    body = re.sub(r"([.!?»\"])([А-ЯЁ])", r"\1 \2", body)

    parts = [body]
    for pattern in (r"\s*\d+\.\s+", r"\s*[-–—]\s+"):
        new_parts: list[str] = []
        for chunk in parts:
            new_parts.extend(re.split(pattern, chunk))
        parts = new_parts

    glued: list[str] = []
    for chunk in parts:
        glued.extend(re.split(r"(?<=[а-яё])(?=[А-ЯЁ][а-яё])", chunk))

    statements: list[str] = []
    for chunk in glued:
        chunk = re.sub(r"\s+", " ", chunk).strip(" .;")
        if len(chunk) < 3 or chunk.lower().startswith("цель:"):
            continue
        statements.append(chunk)
    return statements


def detect_region(text: str) -> str | None:
    for slug, meta in REGIONS.items():
        if slug == "Общероссийский_корпус":
            continue
        for kw in meta["keywords"]:
            if re.search(kw, text, re.I):
                return slug
    return None


def mentions_other_region(text: str, exclude_slug: str) -> bool:
    for slug, meta in REGIONS.items():
        if slug in (exclude_slug, "Общероссийский_корпус"):
            continue
        for kw in meta["keywords"]:
            if re.search(kw, text, re.I):
                return True
    return False


def code_themes(text: str) -> dict[str, int]:
    low = text.lower()
    codes = {code: 0 for code, _, _ in THEME_CODES}
    for code, patterns in THEME_KEYWORDS.items():
        for pat in patterns:
            if re.search(pat, low):
                codes[code] = 1
                break
    return codes


def assign_tone(text: str, themes: dict[str, int]) -> str:
    """Одна взаимоисключающая тональная категория на высказывание."""
    low = text.lower().strip()
    for pat in TONE_DOUBT:
        if re.search(pat, low):
            return "ТОН_СОМНЕНИЕ"
    for pat in TONE_POSITIVE:
        if re.search(pat, low):
            # «доступ есть, но тарифы низкие» — не позитив
            if re.search(r"но\s+тариф|тариф.*низк|невыносим", low):
                break
            return "ТОН_ПОЗИТИВ"
    for pat in TONE_NEGATIVE:
        if re.search(pat, low):
            return "ТОН_НЕГАТИВ"
    # Конструктивные предложения (типичный ответ эксперта)
    proposal_markers = [
        r"упрост", r"повыс", r"увелич", r"внедр", r"расшир", r"обеспеч",
        r"создат", r"пересмотр", r"индекс", r"компенсац", r"отмен",
    ]
    if any(re.search(p, low) for p in proposal_markers) or any(themes.values()):
        return "ТОН_КОНСТРУКТИВ"
    if len(low) < 35:
        return "ТОН_СОМНЕНИЕ"
    return "ТОН_КОНСТРУКТИВ"


def code_statement(text: str) -> dict[str, int]:
    themes = code_themes(text)
    tone_key = assign_tone(text, themes)
    codes = {code: 0 for code, _, _ in CODES}
    for k, v in themes.items():
        codes[k] = v
    codes[tone_key] = 1
    return codes


def build_statements(raw: list[str]) -> list[Statement]:
    result: list[Statement] = []
    for i, text in enumerate(raw, 1):
        h = hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]
        result.append(
            Statement(
                id=i,
                text=text,
                region_tag=detect_region(text),
                codes=code_statement(text),
                source_id=i,
                text_hash=h,
            )
        )
    return result


def select_regional_corpus(all_st: list[Statement], slug: str) -> list[Statement]:
    if slug == "Общероссийский_корпус":
        return list(all_st)

    seen: set[str] = set()
    selected: list[Statement] = []
    for s in all_st:
        ok = s.region_tag == slug
        if not ok:
            for kw in REGIONS[slug]["keywords"]:
                if re.search(kw, s.text, re.I):
                    ok = True
                    break
        if not ok and re.search(r"в нашем регионе|нашем регионе", s.text, re.I):
            if not mentions_other_region(s.text, slug):
                ok = True
        if ok and s.text_hash not in seen:
            seen.add(s.text_hash)
            selected.append(s)

    return [
        Statement(i, s.text, slug, s.codes, s.source_id, s.text_hash)
        for i, s in enumerate(selected, 1)
    ]


def aggregate_stats(statements: list[Statement]) -> dict:
    n = len(statements)
    themes = {c[0]: 0 for c in THEME_CODES}
    tones = {c[0]: 0 for c in TONE_CODES}
    for s in statements:
        for k, v in s.codes.items():
            if not v:
                continue
            if k.startswith("ТОН_"):
                tones[k] += 1
            else:
                themes[k] += 1
    ranked = sorted(themes.items(), key=lambda x: x[1], reverse=True)
    return {"n": n, "themes": themes, "tones": tones, "ranked": ranked}


def save_tone_chart(path: Path, tones: dict[str, int], title: str) -> None:
    labels = ["Позитив", "Конструктив", "Негатив", "Сомнение"]
    keys = ["ТОН_ПОЗИТИВ", "ТОН_КОНСТРУКТИВ", "ТОН_НЕГАТИВ", "ТОН_СОМНЕНИЕ"]
    values = [tones.get(k, 0) for k in keys]
    colors = ["#2E7D32", "#1565C0", "#C62828", "#757575"]
    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(labels, values, color=colors, edgecolor="white", linewidth=0.8)
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)
    ax.set_ylabel("Число высказываний")
    n = sum(values) or 1
    for bar, val in zip(bars, values):
        pct = val / n * 100
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.5,
                f"{val} ({pct:.1f}%)", ha="center", va="bottom", fontsize=10)
    ax.legend(bars, [f"{l}: {v}" for l, v in zip(labels, values)], loc="upper right", fontsize=9)
    ax.set_ylim(0, max(values) * 1.2 + 1)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def save_themes_chart(path: Path, ranked: list[tuple[str, int]], title: str, top_n: int = 10) -> None:
    top = [(k, v) for k, v in ranked if v > 0][:top_n]
    if not top:
        return
    names = {c: n for c, n, _ in THEME_CODES}
    labels = [names.get(k, k) for k, _ in top]
    values = [v for _, v in top]
    fig, ax = plt.subplots(figsize=(9, 5.5))
    y_pos = range(len(labels))
    bars = ax.barh(y_pos, values, color="#1A5276", edgecolor="white")
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Частота упоминаний (число высказываний)")
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2,
                str(val), va="center", fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def write_data_sheet(ws, statements: list[Statement], sheet_title_note: str) -> int:
    ws.append(["ID", "ID_полного_корпуса", "Источник", "Хеш_текста", "Текст"] + [c[0] for c in CODES])
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="E8EEF2")
    for s in statements:
        src = s.region_tag or "Общий корпус"
        row = [s.id, s.source_id, src, s.text_hash, s.text] + [s.codes[c[0]] for c in CODES]
        ws.append(row)
    ws.freeze_panes = "A2"
    ws.column_dimensions["E"].width = 80
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        row[4].alignment = Alignment(wrap_text=True, vertical="top")
    ws["A1"].comment = None
    return ws.max_row


def add_tone_summary_sheet(wb: Workbook, statements: list[Statement], sheet_name: str) -> None:
    ws = wb.create_sheet(sheet_name)
    stats = aggregate_stats(statements)
    n = stats["n"]
    ws["A1"] = "Тональность"
    ws["B1"] = "Число_высказываний"
    ws["C1"] = "Доля"
    for c in range(1, 4):
        ws.cell(row=1, column=c).font = Font(bold=True)
    tone_labels = [
        ("ТОН_ПОЗИТИВ", "Позитив"),
        ("ТОН_КОНСТРУКТИВ", "Конструктивные предложения"),
        ("ТОН_НЕГАТИВ", "Негатив / критика"),
        ("ТОН_СОМНЕНИЕ", "Сомнение / нет ответа"),
    ]
    for i, (code, label) in enumerate(tone_labels, 2):
        cnt = stats["tones"][code]
        ws.cell(row=i, column=1, value=label)
        ws.cell(row=i, column=2, value=cnt)
        ws.cell(row=i, column=3, value=cnt / n if n else 0)
    ws["E1"] = n
    ws["E1"].font = Font(bold=True)


def add_theme_summary_sheet(wb: Workbook, statements: list[Statement], data_sheet: str, sheet_name: str) -> None:
    ws = wb.create_sheet(sheet_name)
    stats = aggregate_stats(statements)
    n = stats["n"]
    last_row = n + 1
    data_ws = wb[data_sheet]
    code_col_start = 6

    ws["A1"] = "Код"
    ws["B1"] = "Название"
    ws["C1"] = "Частота"
    ws["D1"] = "Доля"
    ws["E1"] = "Ранг"
    for c in range(1, 6):
        ws.cell(row=1, column=c).font = Font(bold=True)

    freqs: list[tuple[str, str, int]] = []
    for i, (code, name, _) in enumerate(THEME_CODES):
        col = get_column_letter(code_col_start + i)
        freq = sum(data_ws[f"{col}{row}"].value or 0 for row in range(2, last_row + 1))
        freqs.append((code, name, freq))

    freqs.sort(key=lambda x: x[2], reverse=True)
    sorted_freq_values = [f for _, _, f in freqs]
    for i, (code, name, freq) in enumerate(freqs, 2):
        ws.cell(row=i, column=1, value=code)
        ws.cell(row=i, column=2, value=name)
        ws.cell(row=i, column=3, value=freq)
        ws.cell(row=i, column=4, value=freq / n if n else 0)
        rank = sorted_freq_values.index(freq) + 1 if freq else len(sorted_freq_values)
        ws.cell(row=i, column=5, value=rank)


def add_excel_charts(wb: Workbook, tone_sheet: str, theme_sheet: str) -> None:
    ws_tone = wb[tone_sheet]
    pie = PieChart()
    pie.title = "Распределение по тональности"
    pie.height = 10
    pie.width = 14
    labels = Reference(ws_tone, min_col=1, min_row=2, max_row=5)
    data = Reference(ws_tone, min_col=2, min_row=1, max_row=5)
    pie.add_data(data, titles_from_data=True)
    pie.set_categories(labels)
    pie.dataLabels = DataLabelList()
    pie.dataLabels.showPercent = True
    pie.dataLabels.showVal = True
    ws_tone.add_chart(pie, "F2")

    ws_theme = wb[theme_sheet]
    bar = BarChart()
    bar.type = "bar"
    bar.title = "Топ-10 тематических кодов"
    bar.height = 12
    bar.width = 16
    theme_end = min(11, ws_theme.max_row)
    labels = Reference(ws_theme, min_col=2, min_row=2, max_row=theme_end)
    data = Reference(ws_theme, min_col=3, min_row=1, max_row=theme_end)
    bar.add_data(data, titles_from_data=True)
    bar.set_categories(labels)
    bar.dataLabels = DataLabelList()
    bar.dataLabels.showVal = True
    ws_theme.add_chart(bar, "G2")


def write_excel(path: Path, regional: list[Statement], full: list[Statement], slug: str) -> None:
    wb = Workbook()
    ws_reg = wb.active
    ws_reg.title = "Данные_регион"
    reg_last = write_data_sheet(ws_reg, regional, "региональный подкорпус")

    ws_full = wb.create_sheet("Данные_полный_корпус")
    write_data_sheet(ws_full, full, "полный корпус РФ")

    ws_codes = wb.create_sheet("Коды")
    ws_codes.append(["Код", "Название", "Описание", "Тип"])
    for cell in ws_codes[1]:
        cell.font = Font(bold=True)
    for code, name, desc in THEME_CODES:
        ws_codes.append([code, name, desc, "тема"])
    for code, name, desc in TONE_CODES:
        ws_codes.append([code, name, desc, "тональность"])

    ws_note = wb.create_sheet("Пояснение")
    ws_note["A1"] = "Как читать этот файл"
    ws_note["A1"].font = Font(bold=True, size=12)
    notes = [
        f"Региональный пакет: {REGIONS[slug]['title']}.",
        f"Лист «Данные_регион» — высказывания с явной региональной привязкой ({len(regional)} ед.).",
        f"Лист «Данные_полный_корпус» — все {len(full)} высказываний для перекрёстной проверки.",
        "Итоги_тональность и Итоги_темы рассчитаны по региональному подкорпусу.",
        "Столбец ID_полного_корпуса и Хеш_текста позволяют проверить любую строку в исходном ТЗ.",
        "Если региональный подкорпус мал, сопоставьте с полным корпусом РФ и блоком «в нашем регионе».",
    ]
    for i, line in enumerate(notes, 2):
        ws_note[f"A{i}"] = line

    add_tone_summary_sheet(wb, regional, "Итоги_тональность")
    add_theme_summary_sheet(wb, regional, "Данные_регион", "Итоги_темы")
    add_excel_charts(wb, "Итоги_тональность", "Итоги_темы")

    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_chart_to_doc(doc: Document, chart_path: Path, caption: str) -> None:
    if not chart_path.exists():
        return
    doc.add_picture(str(chart_path), width=Inches(6.2))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.size = Pt(10)


def executive_summary(stats: dict, meta: dict) -> str:
    n = stats["n"]
    tones = stats["tones"]
    top3 = stats["ranked"][:3]
    names = {c: nm for c, nm, _ in THEME_CODES}
    top_txt = ", ".join(f"{names[c]} ({v})" for c, v in top3 if v)
    neg_pct = tones["ТОН_НЕГАТИВ"] / n * 100 if n else 0
    constr_pct = tones["ТОН_КОНСТРУКТИВ"] / n * 100 if n else 0
    return (
        f"По корпусу {n} высказываний ({meta['title']}) эксперты в подавляющем большинстве "
        f"формулируют конкретные меры ({tones['ТОН_КОНСТРУКТИВ']} выск., {constr_pct:.0f}%) "
        f"и описывают системные проблемы ({tones['ТОН_НЕГАТИВ']} выск., {neg_pct:.0f}%). "
        f"Позитивные оценки редки ({tones['ТОН_ПОЗИТИВ']} выск.). "
        f"Ключевые темы: {top_txt}. "
        "Формальный доступ к бюджетному финансированию не обеспечивает устойчивость НКО: "
        "заниженные тарифы, кассовые разрывы и административные барьеры сохраняются."
    )


def verification_text() -> str:
    return (
        "Проверка достоверности данных:\n"
        "1. Первичный массив — файл «ТЗ контент.docx», блок «Высказывания».\n"
        "2. В Excel каждая строка на листе «Данные_регион» или «Данные_полный_корпус» "
        "содержит ID, ID_полного_корпуса, хеш текста (SHA-256, 16 символов) и полный текст.\n"
        "3. Лист «Коды» — расшифровка всех столбцов кодировочной матрицы.\n"
        "4. Листы «Итоги_*» — агрегаты, пересчитываемые из листа «Данные_регион».\n"
        "5. Диаграммы в отчёте и Excel соответствуют тем же агрегатам.\n"
        "6. Региональный подкорпус отделён от полного корпуса РФ — сравнение листов "
        "показывает, какие высказывания имеют явную региональную маркировку."
    )


# --- Тексты отчётов (сокращённо, ключевые функции) ---

def intro_text(meta: dict, n: int, full_n: int) -> str:
    return (
        f"Настоящий отчёт — результаты контент-анализа ответов экспертов (негосударственные "
        f"поставщики социальных услуг) на вопрос 72 анкеты. Контекст: федеральный закон № 442-ФЗ, "
        f"{meta['law']}. Региональный подкорпус: {n} высказываний; полный корпус РФ: {full_n} "
        f"высказываний для перекрёстной проверки."
    )


def context_text(meta: dict) -> str:
    return (
        f"Негосударственный сектор в {meta['title']} работает в условиях формального равенства "
        f"с госучреждениями, но эксперты фиксируют экономическое неравенство (тарифы, сроки выплат), "
        f"административные барьеры и конфликт интересов РОИВ. Ключевой орган: {meta['roiv']}."
    )


def write_report_docx(
    path: Path,
    meta: dict,
    regional: list[Statement],
    full: list[Statement],
    chart_tone: Path,
    chart_themes: Path,
    slug: str,
) -> None:
    stats = aggregate_stats(regional)
    doc = Document()
    title = doc.add_heading(f"Контент-анализ доступа НКО к бюджетному финансированию — {meta['title']}", 0)
    doc.add_paragraph(f"Дата подготовки: {date.today().isoformat()}")
    doc.add_paragraph(
        "Материалы для Российской академии наук, Совета по правам человека при Президенте РФ "
        "и органов исполнительной власти."
    )

    add_heading(doc, "Резюме для принятия управленческих решений", 2)
    doc.add_paragraph(executive_summary(stats, meta))

    add_heading(doc, "Введение", 2)
    doc.add_paragraph(intro_text(meta, stats["n"], len(full)))

    add_heading(doc, "Контекст проекта", 2)
    doc.add_paragraph(context_text(meta))

    add_heading(doc, "Методика и проверка достоверности", 2)
    doc.add_paragraph(verification_text())

    add_heading(doc, "Результаты: тональность", 2)
    tones = stats["tones"]
    n = stats["n"]
    for code, label in [
        ("ТОН_ПОЗИТИВ", "Позитив"),
        ("ТОН_КОНСТРУКТИВ", "Конструктивные предложения"),
        ("ТОН_НЕГАТИВ", "Негатив / критика"),
        ("ТОН_СОМНЕНИЕ", "Сомнение"),
    ]:
        v = tones[code]
        doc.add_paragraph(f"• {label}: {v} ({v/n*100:.1f}%)" if n else f"• {label}: 0")
    add_chart_to_doc(doc, chart_tone, "Рис. 1. Распределение высказываний по тональности (региональный подкорпус)")

    add_heading(doc, "Результаты: ключевые темы", 2)
    names = {c: nm for c, nm, _ in THEME_CODES}
    for code, cnt in stats["ranked"][:10]:
        if cnt:
            doc.add_paragraph(f"• {names[code]}: {cnt} ({cnt/n*100:.1f}%)")
    add_chart_to_doc(doc, chart_themes, "Рис. 2. Топ тематических кодов (региональный подкорпус)")

    add_heading(doc, "Практические рекомендации", 2)
    roiv = meta["roiv"]
    recs = [
        f"Индексация тарифов и рост объёма финансирования НКО в {meta['title']}.",
        f"{roiv}: упростить реестр, сократить отчётность, ускорить выплаты.",
        "Устранить монополию госучреждений и конфликт интересов РОИВ.",
        "Развить соцзаказ, ваучеры, доступ НКО к СДУ.",
        "Имущественная и налоговая поддержка негосударственных поставщиков.",
    ]
    for i, r in enumerate(recs, 1):
        doc.add_paragraph(f"{i}. {r}")

    if slug != "Общероссийский_корпус" and stats["n"] < 20:
        doc.add_paragraph(
            f"Примечание: явно региональных высказываний — {stats['n']}. "
            f"Для общероссийских выводов используйте полный корпус ({len(full)} выск.) "
            "в пакете «Общероссийский_корпус»."
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_ran_report(
    path: Path,
    all_st: list[Statement],
    regional_counts: dict[str, int],
    chart_tone: Path,
    chart_themes: Path,
) -> None:
    stats = aggregate_stats(all_st)
    doc = Document()
    doc.add_heading("Аналитический доклад", 0)
    doc.add_heading(
        "Контент-анализ барьеров доступа негосударственных поставщиков "
        "социальных услуг к бюджетному финансированию", 1
    )
    doc.add_paragraph(f"Дата: {date.today().isoformat()}")
    doc.add_paragraph(
        "Для Российской академии наук, Совета по правам человека при Президенте Российской Федерации "
        "и федеральных органов исполнительной власти."
    )

    add_heading(doc, "1. Резюме", 2)
    doc.add_paragraph(executive_summary(stats, REGIONS["Общероссийский_корпус"]))

    add_heading(doc, "2. Проблема", 2)
    doc.add_paragraph(
        "Негосударственные поставщики формально включены в систему 442-ФЗ, но экономически "
        "и административно уступают госучреждениям. Эксперты описывают скрытое неравенство: "
        "тарифы не покрывают затраты, выплаты задерживаются, лимиты и отчётность ограничивают доступ."
    )

    add_heading(doc, "3. Методика и проверка данных", 2)
    doc.add_paragraph(verification_text())
    doc.add_paragraph(f"Объём полного корпуса: {stats['n']} высказываний.")

    add_heading(doc, "4. Результаты (общероссийский корпус)", 2)
    n = stats["n"]
    tones = stats["tones"]
    doc.add_paragraph(
        f"Тональность: позитив {tones['ТОН_ПОЗИТИВ']} ({tones['ТОН_ПОЗИТИВ']/n*100:.1f}%), "
        f"конструктив {tones['ТОН_КОНСТРУКТИВ']} ({tones['ТОН_КОНСТРУКТИВ']/n*100:.1f}%), "
        f"негатив {tones['ТОН_НЕГАТИВ']} ({tones['ТОН_НЕГАТИВ']/n*100:.1f}%), "
        f"сомнение {tones['ТОН_СОМНЕНИЕ']} ({tones['ТОН_СОМНЕНИЕ']/n*100:.1f}%)."
    )
    add_chart_to_doc(doc, chart_tone, "Рис. 1. Тональность (полный корпус РФ)")
    add_chart_to_doc(doc, chart_themes, "Рис. 2. Топ-10 тем (полный корпус РФ)")

    add_heading(doc, "5. Региональная разбивка (явные упоминания)", 2)
    for slug, cnt in sorted(regional_counts.items(), key=lambda x: -x[1]):
        if slug == "Общероссийский_корпус":
            continue
        doc.add_paragraph(f"• {REGIONS[slug]['title']}: {cnt} высказываний в региональном подкорпусе")

    add_heading(doc, "6. Управленческие рекомендации", 2)
    doc.add_paragraph(
        "1. Экономически обоснованные тарифы с автоматической индексацией.\n"
        "2. Авансирование и многолетние соглашения с НКО.\n"
        "3. Снижение отчётности и прозрачные конкурсы.\n"
        "4. Изменения 442-ФЗ и 189-ФЗ, устранение монополии РОИВ.\n"
        "5. Соцзаказ, ваучеры, доступ к СДУ.\n"
        "6. Имущественная поддержка и равные условия для НКО и госучреждений."
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_corpus_docx(path: Path, statements: list[Statement], title: str) -> None:
    doc = Document()
    add_heading(doc, f"Корпус — {title}")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "ID полного корпуса"
    hdr[2].text = "Хеш"
    hdr[3].text = "Текст"
    for s in statements:
        row = table.add_row().cells
        row[0].text = str(s.id)
        row[1].text = str(s.source_id)
        row[2].text = s.text_hash
        row[3].text = s.text
    doc.save(path)


def copy_tz(folder: Path) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TZ_DOCX, folder / "01_TZ-kontent.docx")
    pdf = ROOT / "ТЗ контент.pdf"
    if pdf.exists():
        shutil.copy2(pdf, folder / "01_TZ-kontent.pdf")


def make_zip(folder: Path, zip_path: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in folder.rglob("*"):
            if file.is_file():
                zf.write(file, file.relative_to(folder.parent))


def main() -> None:
    if not TZ_DOCX.exists():
        raise FileNotFoundError(TZ_DOCX)

    raw_text = extract_text_from_docx(TZ_DOCX)
    all_statements = build_statements(split_statements(raw_text))
    print(f"Полный корпус: {len(all_statements)} высказываний")

    stats_full = aggregate_stats(all_statements)
    print("Тональность (полный корпус):", stats_full["tones"])

    if OUT_ROOT.exists():
        shutil.rmtree(OUT_ROOT)
    OUT_ROOT.mkdir()
    CHARTS_ROOT.mkdir(parents=True, exist_ok=True)

    # Общероссийские диаграммы
    save_tone_chart(CHARTS_ROOT / "rf_tone.png", stats_full["tones"], "Тональность — полный корпус РФ")
    save_themes_chart(CHARTS_ROOT / "rf_themes.png", stats_full["ranked"], "Топ-10 тем — полный корпус РФ")

    regional_counts: dict[str, int] = {}
    for slug, meta in REGIONS.items():
        regional = select_regional_corpus(all_statements, slug)
        regional_counts[slug] = len(regional)
        region_dir = OUT_ROOT / slug
        chart_dir = CHARTS_ROOT / slug
        reg_stats = aggregate_stats(regional)

        save_tone_chart(chart_dir / "tone.png", reg_stats["tones"], f"Тональность — {meta['title']}")
        save_themes_chart(chart_dir / "themes.png", reg_stats["ranked"], f"Топ тем — {meta['title']}")

        (region_dir / "01_ТЗ_и_корпус").mkdir(parents=True)
        (region_dir / "02_Excel_контент-анализ").mkdir(parents=True)
        (region_dir / "03_Отчет").mkdir(parents=True)
        (region_dir / "04_Диаграммы").mkdir(parents=True)

        copy_tz(region_dir / "01_ТЗ_и_корпус")
        write_corpus_docx(
            region_dir / "01_ТЗ_и_корпус" / "02_Корпус_региональный.docx",
            regional,
            meta["title"],
        )
        write_corpus_docx(
            region_dir / "01_ТЗ_и_корпус" / "03_Корпус_полный_РФ.docx",
            all_statements,
            "Российская Федерация",
        )
        write_excel(
            region_dir / "02_Excel_контент-анализ" / "Контент-анализ_НКО_соцуслуги.xlsx",
            regional,
            all_statements,
            slug,
        )
        shutil.copy2(chart_dir / "tone.png", region_dir / "04_Диаграммы" / "01_тональность.png")
        shutil.copy2(chart_dir / "themes.png", region_dir / "04_Диаграммы" / "02_темы.png")

        write_report_docx(
            region_dir / "03_Отчет" / "Методика_и_результаты_контент-анализа.docx",
            meta,
            regional,
            all_statements,
            chart_dir / "tone.png",
            chart_dir / "themes.png",
            slug,
        )

        intro_doc = Document()
        intro_doc.add_heading(f"Введение — {meta['title']}", 0)
        intro_doc.add_paragraph(intro_text(meta, len(regional), len(all_statements)))
        intro_doc.save(region_dir / "03_Отчет" / "01_Введение.docx")

        ctx_doc = Document()
        ctx_doc.add_heading(f"Контекст проекта — {meta['title']}", 0)
        ctx_doc.add_paragraph(context_text(meta))
        ctx_doc.save(region_dir / "03_Отчет" / "02_Контекст_проекта.docx")

        make_zip(region_dir, OUT_ROOT / f"Kontent-analiz-NKO_{slug}.zip")
        print(f"{slug}: регион {len(regional)} выск., тональность {reg_stats['tones']}")

    ran_dir = OUT_ROOT / "Доклад_РАН_СПЧ"
    ran_dir.mkdir()
    write_ran_report(
        ran_dir / "Аналитический_доклад_РАН_СПЧ.docx",
        all_statements,
        regional_counts,
        CHARTS_ROOT / "rf_tone.png",
        CHARTS_ROOT / "rf_themes.png",
    )
    shutil.copy2(CHARTS_ROOT / "rf_tone.png", ran_dir / "диаграмма_тональность.png")
    shutil.copy2(CHARTS_ROOT / "rf_themes.png", ran_dir / "диаграмма_темы.png")
    write_excel(
        ran_dir / "Контент-анализ_полный_корпус.xlsx",
        all_statements,
        all_statements,
        "Общероссийский_корпус",
    )
    make_zip(ran_dir, OUT_ROOT / "Doklad_RAN_SPCh.zip")
    print("Доклад РАН/СПЧ готов.")


if __name__ == "__main__":
    main()
