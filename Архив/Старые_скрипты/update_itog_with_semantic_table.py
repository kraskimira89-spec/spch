from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from recount_semantic_codes import MACRO_CODES, build_outputs


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "Итог.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def main() -> None:
    result = build_outputs()
    ranked = result["ranked"]
    total_n = result["n_atomic"]
    totals = {key: cnt for key, cnt in ranked}
    macro_by_key = {macro.key: macro for macro in MACRO_CODES}

    doc = Document(str(REPORT_PATH))

    if any("Укрупненные смысловые коды" in p.text for p in doc.paragraphs):
        doc.save(str(REPORT_PATH))
        return

    intro = doc.add_paragraph()
    intro.add_run(
        "4. Проверочный пересчет по укрупненным смысловым кодам"
    ).bold = True
    if intro.runs:
        intro.runs[0].font.size = Pt(12)

    p1 = doc.add_paragraph(
        "Дополнительно выполнен проверочный пересчет полного корпуса по 7 укрупненным "
        "смысловым направлениям. Такой шаг нужен, чтобы не дробить близкие формулировки "
        "по нескольким узким кодам и увидеть реальные тематические акценты экспертов."
    )
    if p1.runs:
        p1.runs[0].font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.add_run("Таблица-обоснование смысловых кодов").bold = True

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Тематический код (направление изменений)"
    hdr[1].text = "Что входит в код"
    hdr[2].text = "Частота"
    hdr[3].text = "Доля"
    for cell in hdr:
        shade_cell(cell, "D9E2F3")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for key, cnt in ranked:
        macro = macro_by_key[key]
        row = table.add_row().cells
        row[0].text = macro.name
        row[1].text = macro.includes
        row[2].text = str(cnt)
        row[3].text = f"{cnt / total_n * 100:.1f}%"

    note = doc.add_paragraph(
        "По проверочному пересчету на первое место вышел код «Экономическая модель и тарифная политика» "
        f"— {totals['ЭКОНОМИЧЕСКАЯ_МОДЕЛЬ_И_ТАРИФЫ']} из {total_n} атомарных высказываний "
        f"({totals['ЭКОНОМИЧЕСКАЯ_МОДЕЛЬ_И_ТАРИФЫ'] / total_n * 100:.1f}%). "
        "Это объясняется тем, что сюда объединены не только прямые упоминания повышения тарифов, "
        "но и индексация, пересмотр тарифов, подушевой норматив и компенсация реальных затрат.",
    )
    if note.runs:
        note.runs[0].italic = True

    doc.save(str(REPORT_PATH))


if __name__ == "__main__":
    main()
