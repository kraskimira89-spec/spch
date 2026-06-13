# -*- coding: utf-8 -*-
"""Конвертация markdown-файла в Word (.docx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Вот первое задание. Есть в социологии метод контен.md"
DST = ROOT / "Вот первое задание. Есть в социологии метод контен.docx"


def strip_footnotes(text: str) -> str:
    return re.sub(r"\[\^[^\]]+\]", "", text).strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    text = strip_footnotes(text)
    if not text:
        return
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def parse_table_lines(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = strip_footnotes(val)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("<img"):
            i += 1
            continue

        if line.strip() in ("***", "---", "___"):
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(strip_footnotes(line[2:]), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(strip_footnotes(line[3:]), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(strip_footnotes(line[4:]), level=3)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table_lines(table_lines))
            continue

        if line.strip().startswith("- "):
            add_rich_paragraph(doc, line.strip()[2:], style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if m:
            add_rich_paragraph(doc, m.group(2), style="List Number")
            i += 1
            continue

        if line.strip().startswith("# "):
            doc.add_heading(strip_footnotes(line.strip()[2:]), level=1)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_rich_paragraph(doc, line)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Сохранено: {docx_path} ({docx_path.stat().st_size // 1024} КБ)")


if __name__ == "__main__":
    convert(SRC, DST)
