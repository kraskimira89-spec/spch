from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from build_by_tz import (
    TZ_DOCX,
    TZ_PDF,
    build_corpus,
    build_pdf_segmentation_validation,
    extract_text_from_docx,
    write_pdf_validation_report,
)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "02_Пересчет_смысловых_кодов"


@dataclass(frozen=True)
class MacroCode:
    key: str
    name: str
    description: str
    includes: str
    patterns: tuple[str, ...]


MACRO_CODES: tuple[MacroCode, ...] = (
    MacroCode(
        "ЭКОНОМИЧЕСКАЯ_МОДЕЛЬ_И_ТАРИФЫ",
        "Экономическая модель и тарифная политика",
        "Все высказывания о размере тарифа, подушевом нормативе, индексации и экономической обоснованности стоимости услуги.",
        "Повышение/увеличение тарифов, индексация, пересмотр и актуализация тарифов, подушевые нормативы, полная компенсация фактических затрат, себестоимость услуги.",
        (
            r"\bтариф",
            r"подушев",
            r"норматив",
            r"индекс",
            r"инфляц",
            r"мрот",
            r"стоимост[ьи].*услуг",
            r"экономическ.*обосн",
            r"реальн.*затрат",
            r"фактическ.*затрат",
            r"себестоим",
            r"занижен.*тариф",
            r"низк.*тариф|тариф.*низк",
            r"повышен.*тариф|повысить.*тариф|тариф.*повыш",
            r"увеличен.*тариф|увеличить.*тариф|тариф.*увелич",
            r"пересмотр.*тариф|пересмотрет.*тариф",
            r"скорректир.*тариф|коррекц.*тариф",
            r"поднят.*тариф|поднять.*стоимост",
            r"выравнив.*тариф",
            r"увелич.*финанс",
            r"больше финансирован",
            r"дополнительн.*финанс",
            r"реальн.*финансирован",
            r"финансирован.*по числ",
            r"планирован.*региональн.*ден",
            r"закладыва.*деньг",
        ),
    ),
    MacroCode(
        "БЮРОКРАТИЯ_И_ПРОЦЕДУРЫ",
        "Бюрократическая нагрузка и упрощение процедур",
        "Высказывания о том, что доступ тормозят сложные процедуры, реестры, конкурсы и отчетность.",
        "Реестр поставщиков, конкурсные процедуры, закупки, тендеры, лишние документы, отчетность, проверки, сроки оформления, электронный документооборот.",
        (
            r"реестр",
            r"конкурс",
            r"закуп",
            r"тендер",
            r"отбор",
            r"процедур",
            r"бюрократ",
            r"отч[её]тност",
            r"документооборот",
            r"документ",
            r"справ",
            r"провер",
            r"упрост",
            r"сократ.*срок",
            r"электронн.*формат",
            r"44.?фз",
            r"223.?фз",
        ),
    ),
    MacroCode(
        "ИНСТИТУЦИОНАЛЬНОЕ_РАВЕНСТВО",
        "Институциональное равенство и демонополизация",
        "Высказывания о неравных правилах игры, монополии госучреждений и необходимости менять нормативную среду.",
        "442-ФЗ, 189-ФЗ, региональные акты, равные условия, устранение конфликта интересов, демонополизация, конкуренция, недискриминационный доступ.",
        (
            r"442",
            r"189",
            r"законодатель",
            r"региональн.*акт",
            r"постановлен",
            r"нпа",
            r"монопол",
            r"демонопол",
            r"конкурен",
            r"равн.*услов",
            r"конфликт интерес",
            r"дискримин",
            r"ручн.*регулир",
            r"едины.*правил|един.*стандарт",
            r"госучрежден",
            r"подведомствен",
        ),
    ),
    MacroCode(
        "ИМУЩЕСТВЕННАЯ_И_РЕСУРСНАЯ_ПОДДЕРЖКА",
        "Имущественная и ресурсная поддержка",
        "Высказывания о помещениях, аренде, коммунальных расходах, налоговых льготах и иных ресурсах для устойчивости НКО.",
        "Аренда, коммунальные услуги, помещения, нежилой фонд, имущество, гранты, субсидии, налоговые льготы, кадровые и иные ресурсные условия.",
        (
            r"аренд",
            r"коммунал",
            r"помещен",
            r"имуществ",
            r"нежил.*фонд",
            r"грант",
            r"субсид",
            r"налог",
            r"льгот",
            r"ресурсн",
            r"кадров",
            r"проезд",
            r"оборудован",
        ),
    ),
    MacroCode(
        "ФИНАНСОВАЯ_УСТОЙЧИВОСТЬ",
        "Финансовая устойчивость и порядок выплат",
        "Высказывания о достаточности объема бюджетных средств, стабильности финансирования и своевременности выплат.",
        "Объем финансирования, бюджетные ассигнования, авансирование, долгосрочные договоры, кассовые разрывы, задержки выплат, лимиты, квоты, своевременная компенсация.",
        (
            r"объ[её]м.*финанс|финанс.*объ[её]м",
            r"бюджетн.*финанс",
            r"ассигнован",
            r"аванс",
            r"долгосроч",
            r"многолет",
            r"кассов",
            r"задерж",
            r"срок.*выплат",
            r"своевремен.*выплат",
            r"компенсац",
            r"лимит",
            r"квот",
            r"стабильн.*финанс",
            r"гарант.*объ[её]м",
            r"планируем.*финанс",
            r"своевремен.*оплат",
            r"полная.*компенсац",
            r"остаточн.*принцип",
            r"переходящ.*остат",
            r"не хватает.*финанс",
        ),
    ),
    MacroCode(
        "ИНФОРМАЦИЯ_И_ОБУЧЕНИЕ",
        "Информационная открытость и обучение",
        "Высказывания о прозрачности информации, консультациях, методической поддержке и обучении НКО.",
        "Открытость информации о конкурсах и финансировании, консультации, обучение, семинары, методическая помощь, ресурсные центры, диалог с органами власти.",
        (
            r"информ",
            r"прозрач",
            r"открытост",
            r"консультац",
            r"обучен",
            r"семинар",
            r"тренинг",
            r"методическ",
            r"ресурсн.*центр",
            r"диалог",
            r"совещан",
            r"взаимодейств",
            r"партн[её]р",
            r"понятн.*критер",
        ),
    ),
    MacroCode(
        "ЧЕЛОВЕКОЦЕНТРИЧНЫЕ_МЕХАНИЗМЫ",
        "Человекоцентричные механизмы (Сертификаты, СДУ)",
        "Высказывания, где фокус сделан на выборе получателя услуги и персонализированных механизмах финансирования.",
        "Социальные сертификаты, ваучеры, «деньги следуют за получателем», система долговременного ухода, право выбора поставщика, ИППСУ, критерии нуждаемости.",
        (
            r"сертификат",
            r"ваучер",
            r"деньги следуют",
            r"финанс.*за человек",
            r"долговремен",
            r"\bсду\b",
            r"систем.*уход",
            r"право выбор",
            r"выбор поставщ",
            r"иппсу.*право выбор|право выбор.*иппсу",
            r"иппсу.*нуждающ|нуждающ.*иппсу",
            r"иппсу.*получател|получател.*иппсу",
            r"иппсу.*соц услуг|соц услуг.*иппсу",
            r"нуждающ",
            r"признан.*нужда",
            r"получател",
        ),
    ),
)

TARIFF_GROUPS: tuple[tuple[str, str], ...] = (
    ("Повышение тарифов", r"повышени[ея]\s+тариф|повысить\s+тариф|повышение\s+тарифа"),
    ("Увеличение тарифов", r"увеличени[ея]\s+тариф|увеличение\s+стоимости\s+услуг|увеличение\s+тарифн"),
    ("Пересмотр и актуализация тарифов", r"пересмотр(?:еть|а)?\s+тариф|скорректир\w*\s+тариф|коррекц\w*\s+тариф"),
    ("Поднятие и выравнивание тарифов", r"поднят\w*\s+тариф|поднять\s+стоимост|выравнив\w*\s+тариф"),
)


def code_macro(text: str) -> dict[str, int]:
    low = text.lower()
    result: dict[str, int] = {}
    for macro in MACRO_CODES:
        result[macro.key] = 1 if any(re.search(pat, low) for pat in macro.patterns) else 0
    # Само по себе упоминание ИППСУ не означает человекоцентричный механизм:
    # в ряде ответов это лишь контекст бюджетного планирования или тарифов.
    if "иппсу" in low or "ипссу" in low:
        has_human_context = any(
            re.search(
                pat,
                low,
            )
            for pat in (
                r"право выбор",
                r"выбор поставщ",
                r"нуждающ",
                r"получател",
                r"соц услуг",
                r"долговремен",
                r"\bсду\b",
                r"сертификат",
                r"ваучер",
            )
        )
        if not has_human_context:
            result["ЧЕЛОВЕКОЦЕНТРИЧНЫЕ_МЕХАНИЗМЫ"] = 0
    return result


def tariff_group_counts(statements: list[str]) -> list[tuple[str, int]]:
    counts: list[tuple[str, int]] = []
    for label, pattern in TARIFF_GROUPS:
        cnt = sum(1 for text in statements if re.search(pattern, text.lower()))
        counts.append((label, cnt))
    return counts


def examples_for_macro(statements: list[str], key: str, limit: int = 3) -> list[str]:
    out: list[str] = []
    for text in statements:
        if code_macro(text)[key]:
            txt = text.strip()
            if len(txt) > 190:
                txt = txt[:187] + "..."
            out.append(txt)
            if len(out) >= limit:
                break
    return out


def build_outputs() -> dict[str, object]:
    raw = extract_text_from_docx(TZ_DOCX)
    statements_obj, rough = build_corpus(raw)
    texts = [s.text for s in statements_obj]

    rows: list[dict[str, object]] = []
    totals = {macro.key: 0 for macro in MACRO_CODES}
    for idx, text in enumerate(texts, 1):
        codes = code_macro(text)
        for key, value in codes.items():
            totals[key] += value
        rows.append({"id": idx, "text": text, **codes})

    ranked = sorted(totals.items(), key=lambda item: (-item[1], item[0]))
    tariff_counts = tariff_group_counts(texts)

    OUT.mkdir(parents=True, exist_ok=True)
    xlsx_path = OUT / "01_Матрица_смысловых_кодов.xlsx"
    docx_path = OUT / "02_Пояснительная_записка_по_смысловым_кодам.docx"
    txt_path = OUT / "03_Итоги_смысловых_кодов.txt"
    pdf_check_path = OUT / "05_PDF_check_сегментации.txt"

    wb = Workbook()
    ws = wb.active
    ws.title = "Данные"
    headers = ["ID", "Текст"] + [macro.key for macro in MACRO_CODES]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
    for row in rows:
        ws.append([row["id"], row["text"]] + [row[macro.key] for macro in MACRO_CODES])
    ws.column_dimensions["B"].width = 110

    summary = wb.create_sheet("Итоги")
    summary.append(["Код", "Название", "Частота", "Доля", "Что входит"])
    for cell in summary[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9E2F3")
    total_n = len(texts)
    macro_by_key = {macro.key: macro for macro in MACRO_CODES}
    for key, cnt in ranked:
        macro = macro_by_key[key]
        summary.append([key, macro.name, cnt, cnt / total_n if total_n else 0, macro.includes])
    summary.column_dimensions["B"].width = 42
    summary.column_dimensions["E"].width = 120

    tariff_ws = wb.create_sheet("Тарифы_проверка")
    tariff_ws.append(["Группа формулировок", "Частота"])
    for cell in tariff_ws[1]:
        cell.font = Font(bold=True)
    for label, cnt in tariff_counts:
        tariff_ws.append([label, cnt])
    wb.save(xlsx_path)

    doc = Document()
    doc.add_heading("Пересчет полного корпуса по смысловым кодам", 0)
    doc.add_paragraph(
        f"Проанализирован полный корпус без регионального деления: {len(rough)} исходных блоков и {len(texts)} атомарных высказываний."
    )
    doc.add_paragraph(
        "Новый пересчет выполнен не по формальным кодам ТЗ, а по 7 укрупненным смысловым направлениям. "
        "Одно высказывание могло получить несколько смысловых кодов одновременно."
    )

    doc.add_heading("Что входит в каждый тематический код", level=1)
    for macro in MACRO_CODES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{macro.name}. ").bold = True
        p.add_run(macro.includes)

    doc.add_heading("Итоги пересчета", level=1)
    for pos, (key, cnt) in enumerate(ranked, 1):
        macro = macro_by_key[key]
        doc.add_paragraph(
            f"{pos}. {macro.name} — {cnt} высказываний ({cnt / total_n * 100:.1f}%).",
            style="List Number",
        )

    doc.add_heading("Почему раньше тарифов было мало", level=1)
    doc.add_paragraph(
        "В предыдущем пересчете тема тарифов была сжата в один код «ТАРИФЫ_НОРМАТИВЫ». "
        "При этом родственные формулировки частично распадались по другим темам: компенсация затрат, "
        "финансовая устойчивость, объем финансирования, авансирование, индексирование. "
        "В новом пересчете все прямые и близкие по смыслу формулировки объединены в одно направление "
        "«Экономическая модель и тарифная политика»."
    )

    doc.add_heading("Проверка тарифных формулировок", level=1)
    for label, cnt in tariff_counts:
        doc.add_paragraph(f"{label} — {cnt}", style="List Bullet")

    doc.add_heading("Примеры высказываний по кодам", level=1)
    for macro in MACRO_CODES:
        doc.add_paragraph(macro.name, style="List Number")
        for example in examples_for_macro(texts, macro.key):
            doc.add_paragraph(f"«{example}»", style="List Bullet")

    doc.save(docx_path)

    lines = [
        "Пересчет полного корпуса по смысловым кодам",
        f"Исходных блоков: {len(rough)}",
        f"Атомарных высказываний: {len(texts)}",
        "",
        "Итоги:",
    ]
    for pos, (key, cnt) in enumerate(ranked, 1):
        macro = macro_by_key[key]
        lines.append(f"{pos}. {macro.name} — {cnt} ({cnt / total_n * 100:.1f}%)")
    lines.append("")
    lines.append("Проверка тарифных формулировок:")
    for label, cnt in tariff_counts:
        lines.append(f"- {label}: {cnt}")
    txt_path.write_text("\n".join(lines), encoding="utf-8")

    validation = build_pdf_segmentation_validation(TZ_PDF, rough, statements_obj)
    write_pdf_validation_report(pdf_check_path, validation)

    return {
        "n_rough": len(rough),
        "n_atomic": len(texts),
        "ranked": ranked,
        "xlsx_path": xlsx_path,
        "docx_path": docx_path,
        "txt_path": txt_path,
        "pdf_check_path": pdf_check_path,
        "tariff_counts": tariff_counts,
        "pdf_validation": validation,
    }


if __name__ == "__main__":
    result = build_outputs()
    print(f"Исходных блоков: {result['n_rough']}")
    print(f"Атомарных высказываний: {result['n_atomic']}")
    for idx, (key, cnt) in enumerate(result["ranked"], 1):
        print(f"{idx}. {key}: {cnt}")
    for label, cnt in result["tariff_counts"]:
        print(f"Тарифы::{label}={cnt}")
    print(result["xlsx_path"])
    print(result["docx_path"])
