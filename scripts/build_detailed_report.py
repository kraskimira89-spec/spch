# -*- coding: utf-8 -*-
"""Подробный цветной Word-отчёт по результатам контент-анализа."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_by_tz import (
    OUT,
    THEME_CODES,
    TONE_CODES,
    Statement,
    aggregate,
    build_corpus,
    extract_text_from_docx,
    TZ_DOCX,
)

REPORT_DIR = OUT / "05_Отчет"
CHARTS_DIR = OUT / "04_Диаграммы"

# Палитра
C_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
C_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
C_GREEN = RGBColor(0x54, 0x82, 0x35)
C_RED = RGBColor(0xC0, 0x00, 0x00)
C_ORANGE = RGBColor(0xED, 0x7D, 0x31)
C_GRAY = RGBColor(0x70, 0x70, 0x70)

FILL_HEADER = "1F4E79"
FILL_TOP3 = "D6E4F0"
FILL_MID = "E2EFDA"
FILL_BOX = "FFF2CC"
FILL_TONE_POS = "E2EFDA"
FILL_TONE_NEG = "FCE4D6"
FILL_TONE_DOUBT = "EDEDED"

BAR_COLORS = [
    "#1F4E79", "#2E75B6", "#4472C4", "#5B9BD5", "#70AD47",
    "#548235", "#BF8F00", "#ED7D31", "#C55A11", "#7030A0",
]


def shade_cell(cell, hex_color: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), hex_color)
    el.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(el)


def style_run(run, *, bold=False, size=11, color: RGBColor | None = None) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    style_run(r, bold=True, size=22, color=C_PRIMARY)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(subtitle)
    style_run(r2, size=13, color=C_ACCENT)

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(d.add_run(f"Дата подготовки: {date.today().strftime('%d.%m.%Y')}"), size=10, color=C_GRAY)
    doc.add_paragraph()


def add_info_box(doc: Document, title: str, text: str, fill: str = FILL_BOX) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    style_run(p.add_run(title + "\n"), bold=True, size=12, color=C_PRIMARY)
    style_run(p.add_run(text), size=11)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = C_PRIMARY if level == 1 else C_ACCENT


def add_styled_table_header(row, headers: list[str], fill: str = FILL_HEADER) -> None:
    for i, h in enumerate(headers):
        row.cells[i].text = ""
        p = row.cells[i].paragraphs[0]
        style_run(p.add_run(h), bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(row.cells[i], fill)


def rank_fill(rank: int) -> str:
    if rank <= 3:
        return FILL_TOP3
    if rank <= 6:
        return FILL_MID
    return "FFFFFF"


def example_quotes(statements: list[Statement], code: str, limit: int = 3) -> list[str]:
    out: list[str] = []
    for s in statements:
        if s.codes.get(code) == 1:
            txt = s.text.strip()
            if len(txt) > 220:
                txt = txt[:217] + "…"
            out.append(f"«{txt}»")
            if len(out) >= limit:
                break
    return out


def save_color_ranking_chart(path: Path, ranked: list[tuple[str, int]]) -> None:
    names = {c: n for c, n, _ in THEME_CODES}
    items = [(names[c], v) for c, v in ranked if v > 0][:10]
    if not items:
        return
    labels, values = zip(*items)
    fig, ax = plt.subplots(figsize=(10, 6))
    colors = BAR_COLORS[: len(labels)]
    bars = ax.barh(range(len(labels)), values, color=colors, edgecolor="white", linewidth=0.8)
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels, fontsize=11)
    ax.invert_yaxis()
    ax.set_xlabel("Число высказываний с кодом = 1", fontsize=11)
    ax.set_title("Ранжирование тематических кодов по упоминаемости", fontsize=14, fontweight="bold", color="#1F4E79")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    n = sum(values) or 1
    for bar, val in zip(bars, values):
        pct = val / n * 100
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                f"  {val}  ({pct:.1f}%)", va="center", fontsize=10, fontweight="bold")
    ax.set_xlim(0, max(values) * 1.25)
    fig.patch.set_facecolor("#FAFAFA")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180, facecolor=fig.get_facecolor())
    plt.close(fig)


def save_color_tone_chart(path: Path, tones: dict[str, int], total: int) -> None:
    labels = ["Позитив", "Негатив", "Сомнение"]
    keys = ["ТОНИЧЕСКИЙ_ПОЗИТИВ", "ТОНИЧЕСКИЙ_НЕГАТИВ", "ТОНИЧЕСКИЙ_СОМНЕНИЕ"]
    values = [tones.get(k, 0) for k in keys]
    colors = ["#548235", "#C00000", "#7F7F7F"]
    explode = (0.04, 0.06, 0.04)

    fig, ax = plt.subplots(figsize=(8, 6))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, colors=colors, explode=explode,
        autopct=lambda p: f"{p:.1f}%\n({int(round(p * total / 100))})",
        startangle=90, textprops={"fontsize": 11},
        wedgeprops={"edgecolor": "white", "linewidth": 2},
    )
    for at in autotexts:
        at.set_fontweight("bold")
    ax.set_title("Тональность высказываний экспертов", fontsize=14, fontweight="bold", color="#1F4E79")
    centre = plt.Circle((0, 0), 0.55, fc="white")
    fig.gca().add_artist(centre)
    fig.text(0.5, 0.48, f"n={total}", ha="center", fontsize=12, fontweight="bold", color="#1F4E79")
    fig.patch.set_facecolor("#FAFAFA")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180, facecolor=fig.get_facecolor())
    plt.close(fig)


def save_combined_chart(path: Path, ranked: list[tuple[str, int]], tones: dict[str, int]) -> None:
    """Две диаграммы на одном листе для отчёта."""
    names = {c: n for c, n, _ in THEME_CODES}
    items = [(names[c], v) for c, v in ranked if v > 0][:5]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    fig.patch.set_facecolor("#FAFAFA")

    if items:
        lbl, val = zip(*items)
        ax1.barh(lbl, val, color=BAR_COLORS[: len(lbl)])
        ax1.invert_yaxis()
        ax1.set_title("Топ-5 тем", fontweight="bold", color="#1F4E79")
        ax1.set_xlabel("Частота")

    labels = ["Позитив", "Негатив", "Сомнение"]
    keys = ["ТОНИЧЕСКИЙ_ПОЗИТИВ", "ТОНИЧЕСКИЙ_НЕГАТИВ", "ТОНИЧЕСКИЙ_СОМНЕНИЕ"]
    vals = [tones.get(k, 0) for k in keys]
    ax2.pie(vals, labels=labels, colors=["#548235", "#C00000", "#7F7F7F"], autopct="%1.1f%%", startangle=90)
    ax2.set_title("Тональность", fontweight="bold", color="#1F4E79")

    fig.suptitle("Контент-анализ: ключевые результаты", fontsize=15, fontweight="bold", color="#1F4E79")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180, facecolor=fig.get_facecolor())
    plt.close(fig)


def write_detailed_report(
    path: Path,
    statements: list[Statement],
    rough_count: int,
    chart_rank: Path,
    chart_tone: Path,
    chart_combo: Path,
) -> None:
    stats = aggregate(statements)
    n = stats["n"]
    names = {c: nm for c, nm, _ in THEME_CODES}
    doc = Document()

    # --- Титул ---
    add_title_block(
        doc,
        "ОТЧЁТ О РЕЗУЛЬТАТАХ КОНТЕНТ-АНАЛИЗА",
        "Доступ негосударственных поставщиков социальных услуг\nк бюджетному финансированию",
    )

    add_info_box(
        doc,
        "📋 Резюме для руководителя",
        f"Проанализировано {n} атомарных высказываний экспертов (из {rough_count} исходных ответов "
        f"на вопрос 72 анкеты). Использовано 10 тематических и 3 тональных кода по ТЗ. "
        f"Лидируют темы: {names[stats['ranked'][0][0]]} ({stats['ranked'][0][1]} упом., "
        f"{stats['ranked'][0][1]/n*100:.1f}%), "
        f"{names[stats['ranked'][1][0]]} ({stats['ranked'][1][1]} упом.), "
        f"{names[stats['ranked'][2][0]]} ({stats['ranked'][2][1]} упом.). "
        "Преобладает критическая оценка действующей системы; позитивные оценки редки.",
        FILL_BOX,
    )
    doc.add_paragraph()

    # --- 1. Введение ---
    add_section_heading(doc, "1. Введение", 1)
    doc.add_paragraph(
        "Настоящий отчёт представляет результаты контент-анализа открытых ответов экспертов — "
        "представителей негосударственных поставщиков социальных услуг — на вопрос: "
        "«Какие меры могли бы реально расширить доступ к бюджетному финансированию?»"
    )
    doc.add_paragraph(
        "Исследование выполнено в логике федерального закона № 442-ФЗ «Об основах социального "
        "обслуживания граждан в Российской Федерации». Цель — выявить смысловые паттерны, "
        "оценочные суждения и приоритетные темы; задачи — систематизация корпуса, кодирование, "
        "подсчёт частот, ранжирование тем, анализ тональности."
    )

    # --- 2. Методика ---
    add_section_heading(doc, "2. Методика", 1)
    doc.add_paragraph("Единица анализа — одно атомарное высказывание (одна законченная мысль). "
                      "Составные ответы экспертов разбиты на отдельные строки; одна строка может "
                      "получить несколько тематических кодов (1/0) и один тональный код.")

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    add_styled_table_header(t.rows[0], ["Параметр", "Значение"])
    rows_data = [
        ("Исходных блоков ответов", str(rough_count)),
        ("Атомарных высказываний", str(n)),
        ("Тематических кодов", "10"),
        ("Тональных кодов", "3"),
        ("Источник данных", "ТЗ контент.docx, блок «Высказывания»"),
        ("Инструмент подсчёта", "Excel, листы «Данные», «Итоги_темы», «Итоги_тональность»"),
    ]
    for a, b in rows_data:
        row = t.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.add_paragraph()
    add_section_heading(doc, "2.1. Справочник тематических кодов (10)", 2)
    tc = doc.add_table(rows=1, cols=4)
    tc.style = "Table Grid"
    add_styled_table_header(tc.rows[0], ["№", "Код", "Название", "Содержание"])
    for i, (code, name, desc) in enumerate(THEME_CODES, 1):
        row = tc.add_row().cells
        row[0].text = str(i)
        row[1].text = code
        row[2].text = name
        row[3].text = desc
        if i % 2 == 0:
            for c in row:
                shade_cell(c, "F2F2F2")

    doc.add_paragraph()
    add_section_heading(doc, "2.2. Тональные коды (3)", 2)
    for code, name, desc in TONE_CODES:
        p = doc.add_paragraph(style="List Bullet")
        style_run(p.add_run(f"{code} — {name}: "), bold=True, color=C_ACCENT)
        p.add_run(desc)

    doc.add_page_break()

    # --- 3. Количественные результаты ---
    add_section_heading(doc, "3. Количественные результаты", 1)

    add_info_box(
        doc,
        "📊 Главный вывод",
        "Эксперты воспринимают проблему прежде всего как экономическую и институциональную: "
        "недостаточные объёмы и тарифы финансирования, нестабильность выплат, лимиты, "
        "бюрократия и неравные условия с госучреждениями.",
        FILL_TOP3,
    )
    doc.add_paragraph()

    add_section_heading(doc, "3.1. Ранжирование тематических кодов", 2)
    doc.add_paragraph(
        "Таблица отсортирована по убыванию частоты. Частота — число высказываний, "
        "в которых тематический код = 1."
    )

    rt = doc.add_table(rows=1, cols=5)
    rt.style = "Table Grid"
    add_styled_table_header(rt.rows[0], ["Ранг", "Код", "Название темы", "Частота", "Доля, %"])
    for rank, (code, freq) in enumerate(stats["ranked"], 1):
        if freq == 0:
            continue
        row = rt.add_row().cells
        row[0].text = str(rank)
        row[1].text = code
        row[2].text = names[code]
        row[3].text = str(freq)
        row[4].text = f"{freq / n * 100:.1f}"
        fill = rank_fill(rank)
        for c in row:
            shade_cell(c, fill)
        if rank <= 3:
            for p in row[0].paragraphs + row[2].paragraphs:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph()
    if chart_rank.exists():
        doc.add_picture(str(chart_rank), width=Inches(6.3))
        cap = doc.add_paragraph("Рис. 1. Ранжирование тематических кодов (частота упоминаний)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].font.color.rgb = C_GRAY

    doc.add_paragraph()
    add_section_heading(doc, "3.2. Тональность высказываний", 2)

    tt = doc.add_table(rows=1, cols=4)
    tt.style = "Table Grid"
    add_styled_table_header(tt.rows[0], ["Тональность", "Код", "Частота", "Доля, %"])
    tone_meta = [
        ("Позитив", "ТОНИЧЕСКИЙ_ПОЗИТИВ", FILL_TONE_POS),
        ("Негатив / критика", "ТОНИЧЕСКИЙ_НЕГАТИВ", FILL_TONE_NEG),
        ("Сомнение", "ТОНИЧЕСКИЙ_СОМНЕНИЕ", FILL_TONE_DOUBT),
    ]
    for label, code, fill in tone_meta:
        cnt = stats["tones"][code]
        row = tt.add_row().cells
        row[0].text = label
        row[1].text = code
        row[2].text = str(cnt)
        row[3].text = f"{cnt / n * 100:.1f}"
        for c in row:
            shade_cell(c, fill)

    doc.add_paragraph()
    if chart_tone.exists():
        doc.add_picture(str(chart_tone), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if chart_combo.exists():
        doc.add_paragraph()
        doc.add_picture(str(chart_combo), width=Inches(6.5))
        cap2 = doc.add_paragraph("Рис. 2. Сводная визуализация: топ-5 тем и тональность")
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap2.runs[0].font.size = Pt(10)
        cap2.runs[0].font.color.rgb = C_GRAY

    doc.add_page_break()

    # --- 4. Качественный анализ ---
    add_section_heading(doc, "4. Качественный анализ и иллюстративные цитаты", 1)
    doc.add_paragraph(
        "Ниже — типичные формулировки по пяти наиболее частым тематическим кодам. "
        "Цитаты приведены без указания авторов."
    )

    for rank, (code, freq) in enumerate(stats["ranked"][:5], 1):
        if not freq:
            continue
        add_section_heading(doc, f"4.{rank}. {names[code]} ({freq} упом., {freq/n*100:.1f}%)", 2)
        quotes = example_quotes(statements, code)
        if quotes:
            for q in quotes:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(q)
                run.italic = True
                run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        else:
            doc.add_paragraph("(Иллюстративные цитаты уточняются при ручной верификации кодов.)")

    add_section_heading(doc, "4.6. Латентные смысловые паттерны", 2)
    patterns = [
        ("Скрытое неравенство", "Формально доступ есть, но тарифы занижены, выплаты задерживаются, "
         "конкурсы устроены в пользу госучреждений."),
        ("Монополия и конфликт интересов", "РОИВ одновременно нормотворец, заказчик, контролёр и поставщик услуг."),
        ("Нестабильность", "Короткие горизонты планирования, кассовые разрывы, непредсказуемые лимиты."),
        ("Административная перегрузка", "Избыточная отчётность, дублирование запросов, непрозрачные процедуры."),
    ]
    for title, desc in patterns:
        p = doc.add_paragraph()
        style_run(p.add_run(f"{title}. "), bold=True, color=C_ACCENT)
        p.add_run(desc)

    doc.add_page_break()

    # --- 5. Рекомендации ---
    add_section_heading(doc, "5. Практические рекомендации", 1)
    recs = [
        ("Финансы", "Обеспечить экономически обоснованные тарифы, индексацию, рост объёма финансирования, "
         "авансирование и многолетние соглашения с НКО."),
        ("Процедуры", "Упростить реестр и конкурсы, сократить отчётность, ускорить выплаты, "
         "повысить прозрачность критериев."),
        ("Нормативка", "Скорректировать 442-ФЗ и региональные акты; устранить монополию и конфликт интересов РОИВ."),
        ("Поддержка", "Развивать ресурсные центры, обучение НКО, информирование о мерах финансирования."),
        ("Права получателей", "Отменить искусственные лимиты, обеспечить реальное право выбора поставщика."),
    ]
    for i, (area, text) in enumerate(recs, 1):
        add_info_box(doc, f"Рекомендация {i}. {area}", text, FILL_MID if i % 2 else FILL_TOP3)
        doc.add_paragraph()

    # --- 6. Проверка ---
    add_section_heading(doc, "6. Проверка достоверности данных", 1)
    doc.add_paragraph(
        "1. Первичный массив — файл «ТЗ контент.docx».\n"
        "2. Корпус атомарных высказываний — «02_Корпус_атомарных_высказываний.docx».\n"
        "3. Кодировочная матрица — Excel, лист «Данные» (столбцы кодов 1/0).\n"
        "4. Частоты и ранги — лист «Итоги_темы» (формулы SUM по столбцам).\n"
        "5. Тональность — лист «Итоги_тональность».\n"
        "6. Настоящий отчёт и диаграммы согласованы с теми же агрегатами."
    )

    add_info_box(
        doc,
        "📁 Связанные файлы пакета",
        "02_Excel/Контент-анализ_НКО_соцуслуги.xlsx · "
        "03_Обобщение/Тематические_коды_и_ранжирование.txt · "
        "04_Диаграммы/*.png",
        "E7E6E6",
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def generate_report_package() -> Path:
    """Сгенерировать отчёт и диаграммы (не пересобирает весь пакет)."""
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    CHARTS_DIR.mkdir(parents=True, exist_ok=True)

    raw = extract_text_from_docx(TZ_DOCX)
    statements, rough = build_corpus(raw)
    stats = aggregate(statements)

    chart_rank = CHARTS_DIR / "отчет_ранжирование_тем.png"
    chart_tone = CHARTS_DIR / "отчет_тональность.png"
    chart_combo = CHARTS_DIR / "отчет_сводка.png"

    save_color_ranking_chart(chart_rank, stats["ranked"])
    save_color_tone_chart(chart_tone, stats["tones"], stats["n"])
    save_combined_chart(chart_combo, stats["ranked"], stats["tones"])

    report_path = REPORT_DIR / "Подробный_отчет_контент-анализ.docx"
    write_detailed_report(report_path, statements, len(rough), chart_rank, chart_tone, chart_combo)
    return report_path


if __name__ == "__main__":
    p = generate_report_package()
    print(f"Отчёт сохранён: {p}")
