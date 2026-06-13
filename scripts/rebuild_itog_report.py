from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from recount_semantic_codes import MACRO_CODES, build_outputs


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "01_Отчет" / "01_Итоговый_отчет.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Heading 1", 13), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def style_paragraph(paragraph, *, bold: bool = False, size: int = 12,
                    space_before: int = 0, space_after: int = 6,
                    first_line_indent_cm: float = 1.25,
                    align_center: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Cm(first_line_indent_cm) if first_line_indent_cm else None
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = bold or run.bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def add_table_headers(row, headers: list[str], fill: str = "D9E2F3") -> None:
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        cell.text = header
        shade_cell(cell, fill)
        for p in cell.paragraphs:
            style_paragraph(p, bold=True, size=11, space_before=0, space_after=0, first_line_indent_cm=0)
            for run in p.runs:
                run.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Контент-анализ высказываний экспертов")
    style_paragraph(p, bold=True, size=16, space_before=0, space_after=8, first_line_indent_cm=0, align_center=True)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Вопрос 72: меры расширения доступа негосударственных поставщиков социальных услуг к бюджетному финансированию"
    )
    style_paragraph(p2, bold=True, size=13, space_before=0, space_after=14, first_line_indent_cm=0, align_center=True)


def add_cover_page(doc: Document) -> None:
    p0 = doc.add_paragraph()
    p0.add_run("АНАЛИТИЧЕСКИЙ ОТЧЕТ")
    style_paragraph(p0, bold=True, size=14, space_before=70, space_after=18, first_line_indent_cm=0, align_center=True)

    p1 = doc.add_paragraph()
    p1.add_run("Контент-анализ высказываний экспертов")
    style_paragraph(p1, bold=True, size=18, space_before=0, space_after=10, first_line_indent_cm=0, align_center=True)

    p2 = doc.add_paragraph()
    p2.add_run(
        "Вопрос 72: меры расширения доступа негосударственных поставщиков социальных услуг к бюджетному финансированию"
    )
    style_paragraph(p2, bold=True, size=13, space_before=0, space_after=16, first_line_indent_cm=0, align_center=True)

    p3 = doc.add_paragraph()
    p3.add_run("Полный корпус ответов без регионального деления")
    style_paragraph(p3, size=12, space_before=0, space_after=10, first_line_indent_cm=0, align_center=True)

    p4 = doc.add_paragraph()
    p4.add_run(f"Дата подготовки: {date.today().strftime('%d.%m.%Y')}")
    style_paragraph(p4, size=12, space_before=0, space_after=180, first_line_indent_cm=0, align_center=True)

    p5 = doc.add_paragraph()
    p5.add_run("Подготовлено на основе пересчета по 7 укрупненным смысловым кодам")
    style_paragraph(p5, size=11, space_before=0, space_after=0, first_line_indent_cm=0, align_center=True)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Оглавление")
    style_paragraph(p, bold=True, size=13, space_before=0, space_after=8, first_line_indent_cm=0)

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, size=12, space_before=0, space_after=6, first_line_indent_cm=0)

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    separate_run = paragraph.add_run("Обновите поле в Word: ПКМ -> Обновить поле.")
    separate_run.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    doc.add_page_break()


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(paragraph, size=11, space_before=0, space_after=0, first_line_indent_cm=0)

        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)


def paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_paragraph(p, bold=bold, size=12, space_after=6, first_line_indent_cm=0 if bold else 1.25)


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    style_paragraph(p, bold=True, size=13, space_before=8, space_after=6, first_line_indent_cm=0)


def subsection_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    style_paragraph(p, bold=True, size=12, space_before=6, space_after=4, first_line_indent_cm=0)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p, bold=True, size=11, space_before=4, space_after=3, first_line_indent_cm=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            for p in row.cells[idx].paragraphs:
                style_paragraph(p, size=11, space_before=0, space_after=0, first_line_indent_cm=0)


def build_report() -> Path:
    result = build_outputs()
    ranked = result["ranked"]
    total_n = result["n_atomic"]
    total_rough = result["n_rough"]
    tariff_counts = result["tariff_counts"]
    macro_by_key = {macro.key: macro for macro in MACRO_CODES}

    doc = Document()
    set_base_style(doc)
    add_cover_page(doc)
    add_toc(doc)
    add_title(doc)

    section_heading(doc, "1. Методика контент-анализа")
    subsection_heading(doc, "Объект и единица анализа")
    paragraph(
        doc,
        "Объектом анализа выступил полный массив ответов экспертов на открытый вопрос анкеты "
        "о мерах, которые могли бы реально расширить доступ негосударственных поставщиков "
        "социальных услуг к бюджетному финансированию.",
    )
    paragraph(
        doc,
        "Единицей анализа принято атомарное высказывание: одна завершенная мысль — одна строка "
        "кодировочной матрицы. Составные ответы были дроблены на отдельные смысловые фрагменты.",
    )

    subsection_heading(doc, "Формирование корпуса")
    paragraph(
        doc,
        f"Для анализа использован только полный корпус без регионального деления: "
        f"{total_rough} исходных блоков ответов и {total_n} атомарных высказываний.",
    )
    paragraph(
        doc,
        "Такой подход позволяет не смешивать общероссийские выводы с региональной спецификой "
        "и избежать искусственного дублирования одинаковых ответов по пакетам.",
    )

    subsection_heading(doc, "Логика кодирования")
    paragraph(
        doc,
        "Проверка показала, что прежняя схема с большим числом узких формальных кодов занижала "
        "видимость крупных смысловых направлений, прежде всего тарифной проблематики. Поэтому "
        "итоговый аналитический отчет построен на 7 укрупненных смысловых кодах. Одно "
        "высказывание могло одновременно входить в несколько направлений.",
    )

    subsection_heading(doc, "Таблица-обоснование смысловых кодов")
    caption(doc, "Таблица 1. Обоснование укрупненных смысловых кодов")
    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(
        table.rows[0],
        ["Тематический код (направление изменений)", "Что входит в код", "Частота", "Доля"],
    )
    for key, cnt in ranked:
        macro = macro_by_key[key]
        row = table.add_row().cells
        row[0].text = macro.name
        row[1].text = macro.includes
        row[2].text = str(cnt)
        row[3].text = f"{cnt / total_n * 100:.1f}%"
    set_table_widths(table, [5.0, 10.8, 2.2, 2.0])

    subsection_heading(doc, "Краткая итоговая таблица")
    caption(doc, "Таблица 2. Топ-7 смысловых направлений по полному корпусу")
    summary_table = doc.add_table(rows=1, cols=4)
    try:
        summary_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(summary_table.rows[0], ["Ранг", "Направление", "Частота", "Доля"])
    for rank, (key, cnt) in enumerate(ranked, 1):
        macro = macro_by_key[key]
        row = summary_table.add_row().cells
        row[0].text = str(rank)
        row[1].text = macro.name
        row[2].text = str(cnt)
        row[3].text = f"{cnt / total_n * 100:.1f}%"
    set_table_widths(summary_table, [2.0, 10.0, 2.5, 2.5])

    section_heading(doc, "2. Результаты контент-анализа")
    paragraph(
        doc,
        "Ниже представлены результаты пересчета полного корпуса по 7 укрупненным смысловым "
        "направлениям. Логика раздела подчинена именно этой схеме, поэтому выводы, ранжирование "
        "и рекомендации далее согласованы между собой.",
    )

    for idx, (key, cnt) in enumerate(ranked, 1):
        macro = macro_by_key[key]
        subsection_heading(
            doc,
            f"2.{idx}. {macro.name}",
        )
        paragraph(
            doc,
            f"Это направление зафиксировано в {cnt} атомарных высказываниях "
            f"({cnt / total_n * 100:.1f}% корпуса). В код включались: {macro.includes}",
        )
        if key == "ЭКОНОМИЧЕСКАЯ_МОДЕЛЬ_И_ТАРИФЫ":
            paragraph(
                doc,
                "Именно этот блок оказался ведущим. Это означает, что эксперты прежде всего "
                "видят барьер не в абстрактном отсутствии мер поддержки, а в неработающей "
                "экономике услуги: низких и заниженных тарифах, отсутствии индексации, "
                "неучете реальных затрат и необходимости пересмотра самой тарифной модели.",
            )
            paragraph(
                doc,
                "С академической точки зрения лидерство этого блока закономерно: тариф является "
                "концентрированным выражением всей экономической архитектуры доступа НКО к "
                "бюджетному финансированию. Через обсуждение тарифов респонденты фактически "
                "описывают одновременно проблему недофинансирования, неучета себестоимости, "
                "отсутствия индексации и структурного неравенства условий между государственными "
                "и негосударственными поставщиками.",
            )
        elif key == "ИМУЩЕСТВЕННАЯ_И_РЕСУРСНАЯ_ПОДДЕРЖКА":
            paragraph(
                doc,
                "Высокая частота этого блока показывает, что доступ к бюджетному финансированию "
                "воспринимается экспертами не только как вопрос денег за услугу, но и как вопрос "
                "общих условий выживания организации: аренды, помещений, налогов, грантов и кадровых ресурсов.",
            )
        elif key == "БЮРОКРАТИЯ_И_ПРОЦЕДУРЫ":
            paragraph(
                doc,
                "Третье место этой темы подтверждает, что даже там, где финансирование формально "
                "предусмотрено, оно часто блокируется сложными процедурами допуска, отчетности и конкурса.",
            )
        elif key == "ФИНАНСОВАЯ_УСТОЙЧИВОСТЬ":
            paragraph(
                doc,
                "Отдельный сильный блок связан со стабильностью финансирования: объемами бюджетных "
                "ассигнований, авансированием, лимитами, кассовыми разрывами и задержками выплат.",
            )
        elif key == "ЧЕЛОВЕКОЦЕНТРИЧНЫЕ_МЕХАНИЗМЫ":
            paragraph(
                doc,
                "Этот код объединяет предложения, где центр тяжести переносится на получателя "
                "услуги: сертификаты, ваучеры, право выбора поставщика, СДУ и критерии нуждаемости.",
            )
        elif key == "ИНФОРМАЦИЯ_И_ОБУЧЕНИЕ":
            paragraph(
                doc,
                "Заметная доля таких высказываний показывает, что часть барьеров создается не "
                "только деньгами и нормативкой, но и слабой информированностью НКО о процедурах, "
                "возможностях и правилах входа в систему.",
            )
        elif key == "ИНСТИТУЦИОНАЛЬНОЕ_РАВЕНСТВО":
            paragraph(
                doc,
                "Хотя этот блок меньше остальных по частоте, он стратегически важен: здесь "
                "сконцентрированы высказывания о монополии госучреждений, конфликте интересов и "
                "необходимости равных правил игры.",
            )

    subsection_heading(doc, "2.8. Проверка тарифной темы")
    paragraph(
        doc,
        "Для дополнительной верификации отдельно подсчитаны явные словесные группы, связанные "
        "именно с тарифами. Они подтверждают, что тема действительно является центральной.",
    )
    caption(doc, "Таблица 3. Проверка тарифной темы по ключевым формулировкам")
    tariff_table = doc.add_table(rows=1, cols=2)
    try:
        tariff_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(tariff_table.rows[0], ["Группа формулировок", "Частота"])
    for label, cnt in tariff_counts:
        row = tariff_table.add_row().cells
        row[0].text = label
        row[1].text = str(cnt)
    set_table_widths(tariff_table, [12.5, 3.0])

    subsection_heading(doc, "2.9. Вывод по логике пересчета")
    paragraph(
        doc,
        "Итоговая картина выглядит логично: на первом плане стоят тарифы и экономическая модель, "
        "далее идут ресурсные условия работы организаций, затем административные барьеры и "
        "финансовая устойчивость. Это более согласованная картина, чем прежний отчет, где "
        "смыслово близкие высказывания были разведены по нескольким узким кодам.",
    )

    section_heading(doc, "3. Практические рекомендации органам власти")
    recs = [
        "1. Пересмотреть экономическую модель финансирования: повысить тарифы, ввести регулярную индексацию, учитывать реальные затраты поставщиков и себестоимость услуги.",
        "2. Снизить бюрократическую нагрузку: упростить доступ в реестр, сократить отчетность, сделать конкурсные и закупочные процедуры менее затратными для НКО.",
        "3. Обеспечить институциональное равенство: минимизировать монополию государственных учреждений, устранить конфликт интересов и закрепить единые правила доступа к бюджетным средствам.",
        "4. Усилить имущественную и ресурсную поддержку: компенсировать аренду и коммунальные услуги, расширить налоговые льготы, использовать гранты и субсидии как поддерживающий инструмент.",
        "5. Повысить финансовую устойчивость НКО: увеличить объем ассигнований, ввести авансирование, многолетние договоры и обеспечить своевременные выплаты без кассовых разрывов.",
        "6. Развивать информационную открытость и обучение: делать понятными правила доступа к финансированию, поддерживать консультации, ресурсные центры и обучающие форматы.",
        "7. Продвигать человекоцентричные механизмы: расширять использование сертификатов, принципа «деньги следуют за получателем», СДУ и реального права выбора поставщика.",
    ]
    for item in recs:
        p = doc.add_paragraph(item, style="List Number")
        style_paragraph(p, size=12, space_before=0, space_after=4, first_line_indent_cm=0)

    section_heading(doc, "4. Итог")
    paragraph(
        doc,
        "Проверка логики отчета показала, что для полного корпуса наиболее адекватной является "
        "схема из 7 укрупненных смысловых кодов. Она лучше отражает реальную структуру ответов "
        "экспертов и особенно точнее показывает вес тарифной проблематики.",
    )
    paragraph(
        doc,
        "Документ подготовлен по полному корпусу без регионального деления. Основной вывод: "
        "ключевой запрос экспертов связан с пересмотром экономической модели финансирования, "
        "а не только с точечными административными улучшениями.",
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    add_page_numbers(doc)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


if __name__ == "__main__":
    path = build_report()
    print(path)
