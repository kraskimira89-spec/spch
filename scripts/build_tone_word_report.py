from __future__ import annotations

from datetime import date
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from build_semantic_tone_report import TONE_META, collect_report_data

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "01_Отчет" / "02_Отчет_по_тональности.docx"
CHART_DIR = ROOT / "output" / "01_Отчет" / "_charts"
PIE_PATH = CHART_DIR / "tone_pie.png"
STACKED_PATH = CHART_DIR / "tone_by_meaning_stacked.png"
PCT_PATH = CHART_DIR / "tone_by_meaning_pct.png"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Heading 1", 13), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def style_paragraph(
    paragraph,
    *,
    bold: bool = False,
    size: int = 12,
    space_before: int = 0,
    space_after: int = 6,
    first_line_indent_cm: float = 1.25,
    align_center: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Cm(first_line_indent_cm) if first_line_indent_cm else None
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = bold or run.bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)


def add_table_headers(row, headers: list[str], fill: str = "D9E2F3") -> None:
    for idx, header in enumerate(headers):
        cell = row.cells[idx]
        cell.text = header
        shade_cell(cell, fill)
        for p in cell.paragraphs:
            style_paragraph(p, bold=True, size=11, space_before=0, space_after=0, first_line_indent_cm=0)
            for run in p.runs:
                run.bold = True


def add_cover_page(doc: Document) -> None:
    p0 = doc.add_paragraph()
    p0.add_run("АНАЛИТИЧЕСКИЙ ОТЧЕТ")
    style_paragraph(p0, bold=True, size=14, space_before=70, space_after=18, first_line_indent_cm=0, align_center=True)

    p1 = doc.add_paragraph()
    p1.add_run("Отчет по тональности и смысловым категориям")
    style_paragraph(p1, bold=True, size=18, space_before=0, space_after=10, first_line_indent_cm=0, align_center=True)

    p2 = doc.add_paragraph()
    p2.add_run(
        "Вопрос 72: меры расширения доступа негосударственных поставщиков социальных услуг к бюджетному финансированию"
    )
    style_paragraph(p2, bold=True, size=13, space_before=0, space_after=16, first_line_indent_cm=0, align_center=True)

    p3 = doc.add_paragraph()
    p3.add_run("Полный корпус ответов без регионального деления")
    style_paragraph(p3, size=12, space_before=0, space_after=10, first_line_indent_cm=0, align_center=True)

    p4 = doc.add_paragraph()
    p4.add_run(f"Дата подготовки: {date.today().strftime('%d.%m.%Y')}")
    style_paragraph(p4, size=12, space_before=0, space_after=180, first_line_indent_cm=0, align_center=True)

    p5 = doc.add_paragraph()
    p5.add_run("Подготовлено на основе пересчета по 7 укрупненным смысловым кодам и 4 тональным категориям")
    style_paragraph(p5, size=11, space_before=0, space_after=0, first_line_indent_cm=0, align_center=True)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Оглавление")
    style_paragraph(p, bold=True, size=13, space_before=0, space_after=8, first_line_indent_cm=0)

    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, size=12, space_before=0, space_after=6, first_line_indent_cm=0)

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    separate_run = paragraph.add_run("Обновите поле в Word: ПКМ -> Обновить поле.")
    separate_run.italic = True

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

    doc.add_page_break()


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(paragraph, size=11, space_before=0, space_after=0, first_line_indent_cm=0)

        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)


def paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p, bold=bold, size=12, space_after=6, first_line_indent_cm=0 if bold else 1.25)


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    style_paragraph(p, bold=True, size=13, space_before=8, space_after=6, first_line_indent_cm=0)


def subsection_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    style_paragraph(p, bold=True, size=12, space_before=6, space_after=4, first_line_indent_cm=0)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    style_paragraph(p, bold=True, size=11, space_before=4, space_after=3, first_line_indent_cm=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            for p in row.cells[idx].paragraphs:
                style_paragraph(p, size=11, space_before=0, space_after=0, first_line_indent_cm=0)


def save_tone_pie_chart(path: Path, tone_totals: dict[str, int]) -> None:
    labels = [label for _, label in TONE_META]
    values = [tone_totals[key] for key, _ in TONE_META]
    colors = ["#2E7D32", "#1565C0", "#C62828", "#757575"]
    fig, ax = plt.subplots(figsize=(7, 5))
    ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    ax.set_title("Распределение высказываний по тональности", fontweight="bold")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_stacked_chart(path: Path, cross_rows: list[dict[str, object]]) -> None:
    labels = [row["name"] for row in cross_rows]
    positives = [row["Позитив"] for row in cross_rows]
    constructive = [row["Конструктив"] for row in cross_rows]
    negative = [row["Негатив"] for row in cross_rows]
    doubt = [row["Сомнение"] for row in cross_rows]

    fig, ax = plt.subplots(figsize=(10, 6))
    y = list(range(len(labels)))
    ax.barh(y, positives, color="#2E7D32", label="Позитив")
    left_2 = positives
    ax.barh(y, constructive, left=left_2, color="#1565C0", label="Конструктив")
    left_3 = [a + b for a, b in zip(positives, constructive)]
    ax.barh(y, negative, left=left_3, color="#C62828", label="Негатив")
    left_4 = [a + b + c for a, b, c in zip(positives, constructive, negative)]
    ax.barh(y, doubt, left=left_4, color="#757575", label="Сомнение")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Число высказываний")
    ax.set_title("Смысловые категории по тональности", fontweight="bold")
    ax.legend(loc="lower right")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_pct_chart(path: Path, pct_rows: list[dict[str, object]]) -> None:
    labels = [row["name"] for row in pct_rows]
    positives = [row["Позитив"] * 100 for row in pct_rows]
    constructive = [row["Конструктив"] * 100 for row in pct_rows]
    negative = [row["Негатив"] * 100 for row in pct_rows]
    doubt = [row["Сомнение"] * 100 for row in pct_rows]

    fig, ax = plt.subplots(figsize=(10, 6))
    y = list(range(len(labels)))
    ax.barh(y, positives, color="#2E7D32", label="Позитив")
    left_2 = positives
    ax.barh(y, constructive, left=left_2, color="#1565C0", label="Конструктив")
    left_3 = [a + b for a, b in zip(positives, constructive)]
    ax.barh(y, negative, left=left_3, color="#C62828", label="Негатив")
    left_4 = [a + b + c for a, b, c in zip(positives, constructive, negative)]
    ax.barh(y, doubt, left=left_4, color="#757575", label="Сомнение")
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Доля внутри категории, %")
    ax.set_title("Структура тональности внутри смысловых категорий", fontweight="bold")
    ax.legend(loc="lower right")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180)
    plt.close(fig)


def build_report() -> Path:
    data = collect_report_data()
    total_n = data["n_atomic"]
    total_rough = data["n_rough"]
    tone_totals = data["tone_totals"]
    cross_rows = data["cross_rows"]
    pct_rows = data["pct_rows"]
    examples_by_tone = data["examples_by_tone"]
    top_by_tone = data["top_by_tone"]

    save_tone_pie_chart(PIE_PATH, tone_totals)
    save_stacked_chart(STACKED_PATH, cross_rows)
    save_pct_chart(PCT_PATH, pct_rows)

    doc = Document()
    set_base_style(doc)
    add_cover_page(doc)
    add_toc(doc)

    section_heading(doc, "1. Методика и логика тонального анализа")
    paragraph(
        doc,
        "Объектом анализа выступил полный корпус ответов экспертов на вопрос о мерах, которые могли бы "
        "расширить доступ негосударственных поставщиков социальных услуг к бюджетному финансированию.",
    )
    paragraph(
        doc,
        f"Для анализа использован полный массив без регионального деления: {total_rough} исходных блоков "
        f"и {total_n} атомарных высказываний. Единицей анализа принято атомарное высказывание: одна завершенная мысль — одна строка.",
    )
    paragraph(
        doc,
        "Каждому высказыванию присваивалась одна тональная категория: позитив, конструктив, негатив или сомнение. "
        "Параллельно высказывание могло относиться к нескольким укрупненным смысловым кодам, что позволило "
        "сопоставить эмоциональную окраску с тематическим содержанием.",
    )

    subsection_heading(doc, "1.1. Правила отнесения к тональным категориям")
    tone_table = doc.add_table(rows=1, cols=3)
    try:
        tone_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(tone_table.rows[0], ["Категория", "Смысл интерпретации", "Частота"])
    tone_desc = {
        "Позитив": "Удовлетворенность действующими мерами, указание на отсутствие проблем или достаточность текущих решений.",
        "Конструктив": "Предложения по изменению правил, процедур, тарифов, механизмов финансирования и поддержки.",
        "Негатив": "Критика существующей системы, описание барьеров, убытков, задержек, дискриминации и неработающих механизмов.",
        "Сомнение": "Неопределенность, отсутствие предложений, затруднение с ответом или предельно короткая нерефлексивная реакция.",
    }
    label_by_key = {key: label for key, label in TONE_META}
    for tone_key, label in TONE_META:
        row = tone_table.add_row().cells
        row[0].text = label
        row[1].text = tone_desc[label]
        row[2].text = str(tone_totals[tone_key])
    set_table_widths(tone_table, [3.0, 11.5, 2.5])

    section_heading(doc, "2. Общая картина тональности")
    paragraph(
        doc,
        "Тональная структура корпуса показывает явное преобладание конструктивного типа ответов. "
        "Большинство экспертов не ограничиваются жалобой или общей оценкой, а формулируют конкретные управленческие предложения.",
    )
    caption(doc, "Таблица 1. Распределение высказываний по тональности")
    summary_table = doc.add_table(rows=1, cols=4)
    try:
        summary_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(summary_table.rows[0], ["Тональность", "Частота", "Доля", "Краткий вывод"])
    conclusions = {
        "Позитив": "Позитивные оценки редки и не формируют основную линию корпуса.",
        "Конструктив": "Доминирующий тип реакции: эксперты предлагают решения и корректировки.",
        "Негатив": "Негатив концентрируется вокруг системных барьеров и неравных условий.",
        "Сомнение": "Часть ответов не содержит развернутой позиции или четкого предложения.",
    }
    for tone_key, label in TONE_META:
        cnt = tone_totals[tone_key]
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = str(cnt)
        row[2].text = f"{cnt / total_n * 100:.1f}%"
        row[3].text = conclusions[label]
    set_table_widths(summary_table, [3.0, 2.5, 2.5, 8.0])

    if PIE_PATH.exists():
        doc.add_picture(str(PIE_PATH), width=Inches(5.8))
        caption(doc, "Рисунок 1. Распределение полного корпуса по 4 тональным категориям")

    subsection_heading(doc, "2.1. Интерпретация общей структуры")
    paragraph(
        doc,
        f"Конструктивные высказывания составляют {tone_totals['ТОН_КОНСТРУКТИВ']} из {total_n} единиц "
        f"({tone_totals['ТОН_КОНСТРУКТИВ'] / total_n * 100:.1f}%). Это означает, что экспертный корпус "
        "в целом носит проектный характер: респонденты описывают не столько настроение, сколько желаемые изменения политики.",
    )
    paragraph(
        doc,
        f"Негативные высказывания составляют {tone_totals['ТОН_НЕГАТИВ']} ({tone_totals['ТОН_НЕГАТИВ'] / total_n * 100:.1f}%). "
        "Их доля меньше конструктивной, но именно через них наиболее отчетливо проявляются зоны конфликта: низкие тарифы, "
        "задержки выплат, монополизация и дискриминационные барьеры.",
    )
    paragraph(
        doc,
        f"Позитивных высказываний всего {tone_totals['ТОН_ПОЗИТИВ']} ({tone_totals['ТОН_ПОЗИТИВ'] / total_n * 100:.1f}%), "
        "что подтверждает: корпус в основном не отражает удовлетворенность системой, а фиксирует запрос на изменение правил.",
    )

    section_heading(doc, "3. Тональность внутри смысловых категорий")
    paragraph(
        doc,
        "Ниже показано, как распределяется тональность внутри семи укрупненных смысловых направлений. "
        "Это позволяет понять, какие темы чаще описываются как проблемные, а какие формулируются в виде конструктивных предложений.",
    )

    caption(doc, "Таблица 2. Смысловые категории по тональности")
    cross_table = doc.add_table(rows=1, cols=6)
    try:
        cross_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(cross_table.rows[0], ["Смысловая категория", "Позитив", "Конструктив", "Негатив", "Сомнение", "Итого"])
    for item in cross_rows:
        row = cross_table.add_row().cells
        row[0].text = item["name"]
        row[1].text = str(item["Позитив"])
        row[2].text = str(item["Конструктив"])
        row[3].text = str(item["Негатив"])
        row[4].text = str(item["Сомнение"])
        row[5].text = str(item["Итого"])
    set_table_widths(cross_table, [7.7, 1.7, 2.1, 1.8, 1.8, 1.8])

    if STACKED_PATH.exists():
        doc.add_picture(str(STACKED_PATH), width=Inches(6.5))
        caption(doc, "Рисунок 2. Абсолютное распределение тональности по смысловым категориям")

    caption(doc, "Таблица 3. Доли тональности внутри каждой смысловой категории")
    pct_table = doc.add_table(rows=1, cols=5)
    try:
        pct_table.style = "Table Grid"
    except KeyError:
        pass
    add_table_headers(pct_table.rows[0], ["Смысловая категория", "Позитив", "Конструктив", "Негатив", "Сомнение"])
    for item in pct_rows:
        row = pct_table.add_row().cells
        row[0].text = item["name"]
        row[1].text = f"{item['Позитив'] * 100:.1f}%"
        row[2].text = f"{item['Конструктив'] * 100:.1f}%"
        row[3].text = f"{item['Негатив'] * 100:.1f}%"
        row[4].text = f"{item['Сомнение'] * 100:.1f}%"
    set_table_widths(pct_table, [8.3, 2.0, 2.2, 2.0, 2.0])

    if PCT_PATH.exists():
        doc.add_picture(str(PCT_PATH), width=Inches(6.5))
        caption(doc, "Рисунок 3. Процентная структура тональности внутри каждой смысловой категории")

    subsection_heading(doc, "3.1. Ключевые выводы по тематико-тональному профилю")
    for tone_key, label in TONE_META:
        leaders = top_by_tone[label]
        leader_text = "; ".join(f"{name} ({cnt})" for name, cnt in leaders) if leaders else "нет выраженных лидеров"
        paragraph(doc, f"{label}: ведущие категории — {leader_text}.")

    paragraph(
        doc,
        "Наибольшая концентрация негативных высказываний наблюдается там, где эксперты описывают системные экономические и институциональные барьеры. "
        "В то же время даже эти темы чаще формулируются в конструктивном ключе: через пересмотр тарифов, упрощение процедур и изменение правил финансирования.",
    )

    section_heading(doc, "4. Примеры высказываний по тональности")
    paragraph(
        doc,
        "Ниже приведены характерные формулировки, показывающие, как именно звучит каждая тональная категория в корпусе ответов.",
    )
    for tone_key, label in TONE_META:
        subsection_heading(doc, f"4.{[k for k, _ in TONE_META].index(tone_key) + 1}. {label}")
        for example in examples_by_tone[tone_key]:
            p = doc.add_paragraph(f"«{example}»", style="List Bullet")
            style_paragraph(p, size=11, space_before=0, space_after=3, first_line_indent_cm=0)

    section_heading(doc, "5. Итог")
    paragraph(
        doc,
        "Тональный профиль полного корпуса показывает доминирование конструктивной модели экспертного ответа. "
        "Это означает, что респонденты в массе своей не ограничиваются оценкой существующих проблем, а формулируют перечень конкретных решений.",
    )
    paragraph(
        doc,
        "Негативная тональность важна как индикатор наиболее болезненных участков системы, однако она не преобладает. "
        "Главная особенность корпуса состоит в сочетании критики барьеров с высокой проектной направленностью предложений.",
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    add_page_numbers(doc)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


if __name__ == "__main__":
    path = build_report()
    print(path)
